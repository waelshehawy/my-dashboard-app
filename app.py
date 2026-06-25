# app.py - النسخة النهائية للإنترنت مع Supabase
import streamlit as st
import pandas as pd
import os
import io
import folium
import json
from streamlit_folium import st_folium
from folium.plugins import MarkerCluster
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta, date
import plotly.graph_objects as go
import plotly.express as px
import psycopg2
from psycopg2.extras import RealDictCursor
from supabase import create_client
from datetime import datetime
from datetime import datetime, timedelta

#============================================================
# إعداد الصفحة
#============================================================
st.set_page_config(
    page_title="PreView Ads ERP - نظام إدارة الإعلانات",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)
#============================================================
# إعدادات Supabase (من متغيرات البيئة)
# ============================================================

def get_connection():
    """اتصال مباشر بـ Supabase PostgreSQL باستخدام st.secrets"""
    return psycopg2.connect(
        host=st.secrets["SUPABASE_HOST"],
        port=st.secrets["SUPABASE_PORT"],
        database=st.secrets["SUPABASE_DB"],
        user=st.secrets["SUPABASE_USER"],
        password=st.secrets["SUPABASE_PASSWORD"],
        sslmode="require",
        connect_timeout=30
    )

# ============================================================
# تهيئة session_state (قبل أي شيء آخر)
# ============================================================

if 'auth' not in st.session_state:
    st.session_state.auth = False
if 'role' not in st.session_state:
    st.session_state.role = None
if 'username' not in st.session_state:
    st.session_state.username = None
if 'user_id' not in st.session_state:
    st.session_state.user_id = None
if 'cart' not in st.session_state:
    st.session_state.cart = {}
if 'booking_cart' not in st.session_state:
    st.session_state.booking_cart = []
if 'selected_company' not in st.session_state:
    st.session_state.selected_company = None
if 'show_company_map' not in st.session_state:
    st.session_state.show_company_map = False
if 'selected_city' not in st.session_state:
    st.session_state.selected_city = None
if 'show_city_details' not in st.session_state:
    st.session_state.show_city_details = False
if 'show_all_cities' not in st.session_state:
    st.session_state.show_all_cities = False
# ============================================================
# دوال المصادقة
# ============================================================

def is_authenticated():
    """التحقق من مصادقة المستخدم"""
    return st.session_state.get('auth', False) and st.session_state.get('user_id') is not None

def get_current_user():
    """الحصول على معلومات المستخدم الحالي"""
    if 'user_id' not in st.session_state:
        return None
    
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute('''
            SELECT id, username, role, full_name, created_at 
            FROM users 
            WHERE id = %s
        ''', (st.session_state.user_id,))
        user = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if user:
            return {
                'id': user[0],
                'username': user[1],
                'role': user[2],
                'full_name': user[3],
                'created_at': user[4]
            }
    except Exception as e:
        st.error(f"❌ خطأ في جلب المستخدم: {e}")
    
    return None

def format_period_for_display(period_name, take_first=False):
    """تحويل اسم الفترة إلى صيغة مفهومة للزبون"""
    if period_name is None:
        return ""
    
    # معالجة الأسماء التي تحتوي على مسافات متعددة (مثل 'تشرين ثاني 15-1')
    parts = period_name.split(' ')
    
    # إذا كان الاسم مكوناً من 3 أجزاء (مثل 'تشرين ثاني 15-1')
    if len(parts) == 3:
        month = f"{parts[0]} {parts[1]}"  # 'تشرين ثاني'
        days = parts[2]                    # '15-1'
    else:
        month = parts[0]
        days = parts[1]
    
    day_parts = days.split('-')
    if len(day_parts) == 2:
        day = day_parts[0] if take_first else day_parts[1]
        return f"{day} {month}"
    
    return period_name

def format_date_range(start_period, end_period):
    """تنسيق نطاق الفترات بشكل مفهوم للزبون"""
    start_formatted = format_period_for_display(start_period)
    end_formatted = format_period_for_display(end_period)
    return f"اعتباراً من {start_formatted} لغاية {end_formatted}"
# ============================================================
# التحسينات البصرية
# ============================================================

ADVANCED_CSS = """
<style>
    @keyframes gradientShift {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }
    
    .stApp {
        background: linear-gradient(-45deg, #ee7752, #e73c7e, #23a6d5, #23d5ab);
        background-size: 400% 400%;
        animation: gradientShift 15s ease infinite;
    }
    
    [data-testid="stSidebar"] {
        background: rgba(26, 26, 46, 0.95) !important;
        backdrop-filter: blur(12px) !important;
        border-right: 1px solid rgba(255,255,255,0.2) !important;
    }
    
    [data-testid="stSidebar"] * {
        color: white !important;
    }
    
    .neumorphic-card {
        background: linear-gradient(145deg, #e6e9f0, #ffffff);
        border-radius: 28px;
        box-shadow: 12px 12px 24px rgba(0,0,0,0.1), -12px -12px 24px rgba(255,255,255,0.7);
        padding: 20px;
        margin: 15px 0;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    
    .neumorphic-card:hover {
        transform: translateY(-8px);
    }
    
    .stat-card-3d {
        background: linear-gradient(135deg, #667eea, #764ba2);
        border-radius: 20px;
        padding: 20px;
        text-align: center;
        color: white;
        transition: transform 0.3s ease;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
    }
    
    .stat-card-3d:hover {
        transform: translateY(-5px) scale(1.02);
    }
    
    .stat-number-3d {
        font-size: 48px;
        font-weight: bold;
        animation: numberPulse 2s ease-in-out infinite;
    }
    
    @keyframes numberPulse {
        0%, 100% { transform: scale(1); }
        50% { transform: scale(1.05); text-shadow: 0 0 20px rgba(255,255,255,0.5); }
    }
    
    .stButton > button {
        background: linear-gradient(45deg, #667eea, #764ba2) !important;
        border: none !important;
        border-radius: 50px !important;
        padding: 10px 24px !important;
        color: white !important;
        font-weight: bold !important;
        transition: all 0.3s ease !important;
        width: 100%;
    }
    
    .stButton > button:hover {
        transform: translateY(-3px) !important;
        box-shadow: 0 10px 20px rgba(102,126,234,0.4) !important;
    }
    
    .badge-animated {
        display: inline-block;
        padding: 6px 16px;
        border-radius: 30px;
        font-size: 12px;
        font-weight: bold;
        margin: 3px;
        animation: badgePop 0.5s ease-out;
    }
    
    @keyframes badgePop {
        from { transform: scale(0); opacity: 0; }
        to { transform: scale(1); opacity: 1; }
    }
    
    [data-testid="stMetric"] {
        background: linear-gradient(135deg, #667eea15, #764ba215);
        border-radius: 16px;
        padding: 16px;
        backdrop-filter: blur(4px);
        transition: all 0.3s ease;
    }
    
    [data-testid="stMetric"]:hover {
        transform: translateY(-5px);
    }
    
    .dataframe {
        border-radius: 15px !important;
        overflow: hidden !important;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1) !important;
    }
    
    .dataframe th {
        background: linear-gradient(135deg, #667eea, #764ba2) !important;
        color: white !important;
        font-weight: bold !important;
        padding: 12px !important;
    }
    
    .dataframe tr:hover {
        background: rgba(102,126,234,0.1) !important;
    }
</style>
"""



st.markdown(ADVANCED_CSS, unsafe_allow_html=True)


# ============================================================
# دوال المتاح
# ============================================================
# ==================== دوال الأعمدة المتاحة (جديدة) ====================

MONTHS_AR = {
    1: "كانون ثاني", 2: "شباط", 3: "اذار", 4: "نيسان",
    5: "ايار", 6: "حزيران", 7: "تموز", 8: "اب",
    9: "ايلول", 10: "تشرين اول", 11: "تشرين ثاني", 12: "كانون اول"
}

def convert_date_to_period_name(date):
    month_name = MONTHS_AR[date.month]
    if date.day <= 15:
        return f"{month_name} 15-1"
    else:
        return f"{month_name} 30-15"

def get_available_boards_from_date(start_date):
    """اللوحات المتاحة ابتداءً من تاريخ محدد - تستخدم اتصال PostgreSQL الموجود"""
    target_period = convert_date_to_period_name(start_date)
    target_year = start_date.year
    
    conn = get_connection()
    cursor = conn.cursor()
    
    # جلب أرقام اللوحات المحجوزة في الفترة المطلوبة
    cursor.execute("""
        SELECT "رقم اللوحة" FROM "حجوزات1" 
        WHERE "فترة الحجز" = %s AND "العام" = %s
    """, (target_period, target_year))
    
    booked_ids = [row[0] for row in cursor.fetchall()]
    
    # جلب جميع الأعمدة
    cursor.execute('SELECT * FROM "اعمدة انارة"')
    all_columns = cursor.fetchall()
    col_names = [desc[0] for desc in cursor.description]
    
    cursor.close()
    conn.close()
    
    all_boards_df = pd.DataFrame(all_columns, columns=col_names)
    available_df = all_boards_df[~all_boards_df['رقم اللوحة'].isin(booked_ids)]
    
    return available_df

# ============================================================
# دوال تحويل الفترات
# ============================================================

PERIOD_ORDER = {
    # كانون الثاني
    'كانون ثاني 15-1': 1, 'كانون ثاني 30-15': 2,
    # شباط
    'شباط 15-1': 3, 'شباط 30-15': 4,
    # آذار
    'اذار 15-1': 5, 'اذار 30-15': 6,
    # نيسان
    'نيسان 15-1': 7, 'نيسان 30-15': 8,
    # أيار
    'ايار15-1': 9, 'أيار 30-15': 10,
    # حزيران
    'حزيران 15-1': 11, 'حزيران 30-15': 12,
    # تموز
    'تموز 15-1': 13, 'تموز 30-15': 14,
    # آب
    'اب 15-1': 15, 'اب 30-15': 16,
    # أيلول
    'أيلول 15-1': 17, 'ايلول30-15': 18,
    # تشرين الأول
    'تشرين اول 15-1': 19, 'تشرين اول30-15': 20,
    # تشرين الثاني
    'تشرين ثاني 15-1': 21, 'تشرين ثاني 30-15': 22,
    # كانون الأول
    'كانون اول 15-1': 23, 'كانون اول 30-15': 24
}
def get_period_number(period_name):
    """تحويل اسم الفترة إلى رقم (1-24)"""
    if period_name is None:
        return 99
    return PERIOD_ORDER.get(period_name, 99)

def get_period_from_date(date_obj):
    day = date_obj.day
    month = date_obj.month
    
    month_names = {
        1: 'كانون ثاني', 2: 'شباط', 3: 'اذار', 4: 'نيسان',
        5: 'ايار', 6: 'حزيران', 7: 'تموز', 8: 'اب',
        9: 'ايلول', 10: 'تشرين اول', 11: 'تشرين ثاني', 12: 'كانون اول'
    }
    
    month_name = month_names[month]
    
    if day <= 15:
        period_name = f"{month_name} 15-1"
    else:
        period_name = f"{month_name} 30-15"
    
    return PERIOD_ORDER.get(period_name, 99)
    

# ============================================================
# دوال مساعدة
# ============================================================

def create_metric_card_3d(title, value, icon, color_gradient="primary"):
    gradients = {
        "primary": "linear-gradient(135deg, #667eea, #764ba2)",
        "success": "linear-gradient(135deg, #11998e, #38ef7d)",
        "danger": "linear-gradient(135deg, #f093fb, #f5576c)",
        "warning": "linear-gradient(135deg, #fa709a, #fee140)"
    }
    
    try:
        if isinstance(value, (int, float)) and not isinstance(value, bool):
            formatted_value = f"{value:,}"
        else:
            formatted_value = str(value)
    except:
        formatted_value = str(value)
    
    return f"""
    <div class="stat-card-3d" style="background: {gradients.get(color_gradient, gradients['primary'])}">
        <div style="font-size: 36px; opacity: 0.8;">{icon}</div>
        <div class="stat-number-3d">{formatted_value}</div>
        <div style="font-size: 14px; opacity: 0.9;">{title}</div>
    </div>
    """

def badge_animated(text, badge_type="info"):
    return f'<span class="badge-animated">{text}</span>'

def safe_split(value):
    if value is None or pd.isna(value):
        return []
    if isinstance(value, float):
        return []
    value_str = str(value)
    if value_str in ['', 'nan', 'None', 'NaN']:
        return []
    return [v.strip() for v in value_str.split(',') if v.strip()]

def is_admin():
    return st.session_state.get('role') == 'admin'

def _force_rtl_style(p):
    # ✅ تغيير من LEFT إلى RIGHT
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    for run in p.runs:
        rPr = run._element.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
def set_table_rtl(table):
    """تحويل اتجاه الجدول إلى RTL مع عكس الخلايا"""
    # 1. عكس اتجاه الجدول
    tblPr = table._element.xpath('w:tblPr')[0]
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)
    
    # 2. عكس اتجاه كل خلية في الجدول
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # تطبيق RTL على كل فقرة في الخلية
                pPr = paragraph._element.get_or_add_pPr()
                bidi_paragraph = OxmlElement('w:bidi')
                bidi_paragraph.set(qn('w:val'), '1')
                pPr.append(bidi_paragraph)
                
                # محاذاة النص لليمين
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                for run in paragraph.runs:
                    rPr = run._element.get_or_add_rPr()
                    rtl = OxmlElement('w:rtl')
                    rtl.set(qn('w:val'), '1')
                    rPr.append(rtl)

SYRIA_COORDS = {
    "دمشق": [33.5138, 36.2765],
    "ريف دمشق": [33.45, 36.35],
    "حلب": [36.2028, 37.1343],
    "حمص": [34.7328, 36.7156],
    "حماة": [35.135, 36.748],
    "اللاذقية": [35.531, 35.79],
    "طرطوس": [34.883, 35.883],
    "سوريا": [34.802, 38.996]
}

# تهيئة حالة الجلسة
if "auth" not in st.session_state:
    st.session_state.auth = False
if "cart" not in st.session_state:
    st.session_state.cart = {}
if "temp_cust" not in st.session_state:
    st.session_state.temp_cust = ""


# ============================================================
# صفحة تسجيل الدخول
# ============================================================

if not st.session_state.auth:
    st.markdown("""
    <div style="display: flex; justify-content: center; align-items: center; min-height: 80vh;">
        <div style="background: rgba(255,255,255,0.1); backdrop-filter: blur(10px); border-radius: 30px; padding: 40px; width: 100%; max-width: 450px; text-align: center; box-shadow: 0 20px 40px rgba(0,0,0,0.2);">
            <div style="width: 80px; height: 80px; background: linear-gradient(135deg, #667eea, #764ba2); border-radius: 50%; display: flex; align-items: center; justify-content: center; margin: 0 auto 20px;">
                <span style="font-size: 40px;">🎯</span>
            </div>
            <h1 style="color: white;">PreView Ads</h1>
            <p style="color: rgba(255,255,255,0.7);">نظام إدارة الإعلانات</p>
    """, unsafe_allow_html=True)
    
    with st.form("login_form"):
        username = st.text_input("👤 اسم المستخدم", placeholder="أدخل اسم المستخدم")
        password = st.text_input("🔑 كلمة المرور", type="password", placeholder="أدخل كلمة المرور")
        submitted = st.form_submit_button("🚪 دخول", use_container_width=True)
        
        if submitted:
            try:
                conn = get_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT id, username, password, role FROM users WHERE username = %s AND password = %s",
                    (username, password)
                )
                user = cursor.fetchone()
                cursor.close()
                conn.close()
                
                if user:
                    st.session_state.auth = True
                    st.session_state.role = user[3]
                    st.session_state.username = user[1]
                    st.session_state.user_id = user[0]
                    st.rerun()
                else:
                    st.error("❌ اسم المستخدم أو كلمة المرور غير صحيحة")
            except Exception as e:
                st.error(f"❌ خطأ في الاتصال: {str(e)}")
    
    st.markdown("</div></div>", unsafe_allow_html=True)
    st.stop()
# ============================================================
# الاتصال بقاعدة البيانات بعد تسجيل الدخول
# ============================================================

conn = get_connection()

# ============================================================
# الشريط الجانبي
# ============================================================


with st.sidebar:
    st.markdown("""
    <div style="text-align: center; padding: 20px 0;">
        <div style="width: 80px; height: 80px; background: linear-gradient(135deg, #667eea, #764ba2); border-radius: 50%; display: flex; align-items: center; justify-content: center; margin: 0 auto;">
            <span style="font-size: 40px;">🎯</span>
        </div>
        <h2 style="color: white; margin-top: 15px;">PreView Ads</h2>
        <p style="color: #a0a0a0; font-size: 12px;">نظام إدارة الإعلانات v2.0</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.divider()

    # معلومات المستخدم المنسقة داخل السايدبار
    user_icon = "👑" if is_admin() else "👤"
    st.markdown(f"""
    <div style="background: rgba(255,255,255,0.1); border-radius: 15px; padding: 15px; text-align: center; margin: 10px 0;">
        <div style="font-size: 30px;">{user_icon}</div>
        <div style="font-weight: bold;">{st.session_state.get('username', '')}</div>
        <div style="font-size: 12px; opacity: 0.7;">{'مدير النظام' if is_admin() else 'موظف'}</div>
    </div>
    """, unsafe_allow_html=True)
    
    # القائمة الرئيسية
    selected_page = st.radio("📋 القائمة الرئيسية", [
        "🏢 لوحات الشركات",
        "📍 الأعمدة المتاحة",
        "📅 لوحة الفترات",
        "📊 Dashboard",
        "📄 عرض سعر",
        "📋 تقرير الجرد",
        "📅 تقرير التوفر الشهري",
        "🗺️ تقرير جميع المواقع",
        "📐 تقرير تجميعي حسب الحجوم",
        "⚙️ الإعدادات",
        "📝 الإدخال اليومي",
        "📋 كتالوج عام"
    ], key="main_menu")
    
    # تحديث الصفحة عند التغيير
    if selected_page != st.session_state.get('page'):
        st.session_state.page = selected_page
        st.rerun()
    
    st.divider()
    
    # زر تسجيل الخروج
    if st.button("🚪 تسجيل الخروج", use_container_width=True):
        st.session_state.auth = False
        st.session_state.cart = {}
        st.rerun()


# ============================================================
# دوال استعلامات Supabase (بصيغة PostgreSQL)
# ============================================================

def run_query(query, params=None, fetch=True):
    """تنفيذ استعلام على Supabase"""
    cursor = conn.cursor()
    try:
        cursor.execute(query, params or ())
        if fetch and query.strip().upper().startswith('SELECT'):
            columns = [desc[0] for desc in cursor.description]
            rows = cursor.fetchall()
            return pd.DataFrame(rows, columns=columns)
        else:
            conn.commit()
            return cursor.rowcount
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        cursor.close()

def get_fees(draw_df, size, print_type, is_foreign):
    subset = draw_df[draw_df['الحجم'] == size].copy()
    
    if print_type == "عادي":
        f_pr = subset[subset['اسم الرسم'].str.contains("اجور الطباعة عادي", na=False)]
        if f_pr.empty:
            f_pr = subset[subset['اسم الرسم'].str.contains("اجور الطباعة", na=False)]
    else:
        f_pr = subset[subset['اسم الرسم'].str.contains("اجور الطباعة", na=False)]
        f_pr = f_pr[~f_pr['اسم الرسم'].str.contains("عادي", na=False)]
    
    fee_print = float(f_pr['اجرة الرسم'].iloc[0]) if not f_pr.empty else 0.0
    
    if is_foreign:
        f_ad = subset[subset['اسم الرسم'].str.contains("اجور العرض اجنبي", na=False)]
        if f_ad.empty:
            f_ad = subset[subset['اسم الرسم'].str.contains("اجور العرض", na=False)]
    else:
        f_ad = subset[subset['اسم الرسم'].str.contains("اجور العرض", na=False)]
        f_ad = f_ad[~f_ad['اسم الرسم'].str.contains("اجنبي", na=False)]
    
    fee_ads = float(f_ad['اجرة الرسم'].iloc[0]) if not f_ad.empty else 0.0
    
    return fee_print, fee_ads

def get_company_bookings():
    """استرجاع بيانات الشركات المحجوزة"""
    query = '''
        SELECT 
            "اسم الزبون" as company_name,
            COUNT(DISTINCT "رقم اللوحة") as total_boards,
            COUNT(DISTINCT "فترة الحجز") as total_periods,
            MAX("العام") as last_year,
            MAX("فترة الحجز") as last_period
        FROM "حجوزات1"
        GROUP BY "اسم الزبون"
        ORDER BY "اسم الزبون"
    '''
    return run_query(query)

def get_company_locations_with_map(company_name):
    """استرجاع مواقع شركة معينة مع الإحداثيات"""
    query = '''
        SELECT DISTINCT 
            b."رقم اللوحة",
            b."اسم العمود",
            b."المحافظة",
            b."الشبكة",
            b."الحجم",
            b."العدد",
            b."Latitude",
            b."Longitude"
        FROM "اعمدة انارة" b
        INNER JOIN "حجوزات1" h ON b."رقم اللوحة" = h."رقم اللوحة"
        WHERE h."اسم الزبون" = %s
    '''
    return run_query(query, (company_name,))

def get_available_by_city():
    """استرجاع الأعمدة المتاحة مجمعة حسب المحافظة"""
    current_year = datetime.now().year
    
    booked_query = 'SELECT DISTINCT "رقم اللوحة" FROM "حجوزات1" WHERE "العام" = %s'
    booked_df = run_query(booked_query, (current_year,))
    booked_boards = booked_df['رقم اللوحة'].tolist() if booked_df is not None and not booked_df.empty else []
    
    all_columns = run_query('SELECT * FROM "اعمدة انارة"')
    
    available = all_columns[~all_columns['رقم اللوحة'].isin(booked_boards)]
    
    def classify_size_for_card(size):
        size_str = str(size).strip()
        if size_str in ['2*1', '2x1', '2 × 1']:
            return 'أعمدة إنارة (2×1)'
        elif size_str in ['125*185', '125x185', '125 × 185']:
            return 'منصفات (125×185)'
        else:
            return 'أحجام أخرى'
    
    available['size_group'] = available['الحجم'].apply(classify_size_for_card)
    return available

def manage_expired_offers():
    """إدارة العروض المنتهية"""
    st.subheader("⚠️ إدارة العروض التي تجاوزت 48 ساعة")
    
    query = '''
        SELECT id, client_name, offer_date 
        FROM "offers_history" 
        WHERE status = 'Pending' AND offer_date < NOW() - INTERVAL '48 hours'
    '''
    expired_df = run_query(query)
    
    if expired_df is None or expired_df.empty:
        st.success("✅ لا توجد عروض منتهية الصلاحية.")
        return
    
    for _, row in expired_df.iterrows():
        col1, col2, col3 = st.columns([3, 1, 1])
        col1.write(f"👤 الزبون: **{row['client_name']}** - تاريخ العرض: {row['offer_date']}")
        
        if is_admin():
            if col2.button("✅ تمديد 48 ساعة", key=f"ext_{row['id']}"):
                cur = conn.cursor()
                cur.execute('UPDATE "offers_history" SET offer_date = NOW() WHERE id = %s', (row['id'],))
                conn.commit()
                cur.close()
                st.success("تم التمديد بنجاح")
                st.rerun()
            
            if col3.button("❌ إلغاء العرض", key=f"del_{row['id']}"):
                cur = conn.cursor()
                cur.execute('UPDATE "offers_history" SET status = %s WHERE id = %s', ('Cancelled', row['id']))
                conn.commit()
                cur.close()
                st.success("تم إلغاء العرض")
                st.rerun()
        else:
            col2.write("🔒")
            col3.write("🔒")

def filter_valid_coordinates(df, lat_col='Latitude', lon_col='Longitude'):
    """تصفية البيانات للحصول على الإحداثيات الصالحة فقط"""
    if df.empty:
        return df
    
    if lat_col not in df.columns or lon_col not in df.columns:
        return pd.DataFrame()
    
    valid = df[
        df[lat_col].notna() & 
        df[lon_col].notna() &
        (df[lat_col] != 0) &
        (df[lon_col] != 0)
    ].copy()
    
    return valid



# ============================================================
#  بداية الصفحات  
# ============================================================
page = st.session_state.get('page', "🏢 لوحات الشركات")

#============================================================
# الصفحة الأولى لوحات الشركات
#============================================================
# ============================================================
# صفحة: لوحات الشركات (باستخدام دوال الفترات)
# ============================================================

if page == "🏢 لوحات الشركات":
    st.title("🏢 لوحات الشركات المعلنة")
    st.markdown('<div class="custom-divider"></div>', unsafe_allow_html=True)
    
    if st.session_state.get('processing', False):
        st.warning("⏳ جاري تحميل البيانات، الرجاء الانتظار...")
        st.stop()
    
    @st.cache_data(ttl=300)
    def load_companies_data():
        try:
            conn = get_connection()
            
            # جلب الفترات الحالية والمستقبلية
            periods_df = pd.read_sql_query('''
                SELECT namee, no FROM "الفترة" 
                WHERE no >= (SELECT MIN(no) FROM "الفترة" WHERE no >= (
                    SELECT no FROM "الفترة" 
                    WHERE namee = (
                        SELECT "فترة الحجز" FROM "حجوزات1" 
                        WHERE "العام" = EXTRACT(YEAR FROM CURRENT_DATE)
                        ORDER BY "فترة الحجز" DESC LIMIT 1
                    )
                ))
                ORDER BY no
            ''', conn)
            
            current_periods = periods_df['namee'].tolist() if not periods_df.empty else []
            
            if current_periods:
                placeholders = ','.join([f"'{p}'" for p in current_periods])
                
                # جلب الحجوزات مع الفترات
                bookings_df = pd.read_sql_query(f'''
                    SELECT 
                        "اسم الزبون" as company_name,
                        COUNT(DISTINCT "رقم اللوحة") as total_boards,
                        STRING_AGG(DISTINCT "فترة الحجز", ' | ') as periods_list,
                        STRING_AGG(DISTINCT CAST("رقم اللوحة" AS TEXT), ', ') as boards_list,
                        MIN("فترة الحجز") as first_period,
                        MAX("فترة الحجز") as last_period,
                        "العام"
                    FROM "حجوزات1" 
                    WHERE "فترة الحجز" IN ({placeholders})
                    AND "العام" >= EXTRACT(YEAR FROM CURRENT_DATE) - 1
                    GROUP BY "اسم الزبون", "العام"
                    ORDER BY "اسم الزبون"
                ''', conn)
                
                # حساب عدد الفترات باستخدام get_period_number()
                if not bookings_df.empty:
                    def calculate_period_count(row):
                        periods = row['periods_list'].split(' | ') if row['periods_list'] else []
                        # استخدام get_period_number() لحساب الفترات الفريدة
                        unique_periods = set()
                        for p in periods:
                            period_num = get_period_number(p.strip())
                            if period_num != 99:
                                unique_periods.add(period_num)
                        return len(unique_periods)
                    
                    bookings_df['total_periods'] = bookings_df.apply(calculate_period_count, axis=1)
                else:
                    bookings_df['total_periods'] = 0
            else:
                bookings_df = pd.DataFrame()
            
            conn.close()
            return {'bookings': bookings_df, 'current_periods': current_periods}
        except Exception as e:
            st.error(f"❌ خطأ في تحميل البيانات: {str(e)}")
            return None
    
    with st.spinner("🔄 جاري تحميل بيانات الشركات..."):
        data = load_companies_data()
    
    if data is None:
        st.error("❌ فشل تحميل البيانات")
        st.stop()
    
    bookings_df = data['bookings']
    current_periods = data['current_periods']
    
    if bookings_df is not None and not bookings_df.empty:
        # حساب المدة وترتيب الشركات
        bookings_df['period_duration'] = bookings_df['total_periods']
        bookings_df = bookings_df.sort_values('period_duration', ascending=False)
        
        total_companies = len(bookings_df)
        total_boards = bookings_df['total_boards'].sum()
        total_periods = bookings_df['total_periods'].sum()
        
        col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
        with col_stat1:
            st.metric("🏢 إجمالي الشركات", total_companies)
        with col_stat2:
            st.metric("📊 إجمالي اللوحات", total_boards)
        with col_stat3:
            st.metric("📅 إجمالي الفترات", total_periods)
        with col_stat4:
            st.metric("🟢 الفترات النشطة", len(current_periods))
        
        st.markdown("---")
        
        col_filter1, col_filter2 = st.columns([2, 1])
        with col_filter1:
            search_term = st.text_input("🔍 بحث عن شركة:", placeholder="اكتب اسم الشركة...")
        with col_filter2:
            sort_by = st.selectbox(
                "ترتيب حسب:",
                ["المدة (الأطول)", "المدة (الأقصر)", "عدد اللوحات (الأكثر)", "عدد اللوحات (الأقل)", "اسم الشركة"]
            )
        
        filtered_df = bookings_df.copy()
        if search_term:
            filtered_df = filtered_df[filtered_df['company_name'].str.contains(search_term, case=False, na=False)]
        
        if sort_by == "المدة (الأطول)":
            filtered_df = filtered_df.sort_values('period_duration', ascending=False)
        elif sort_by == "المدة (الأقصر)":
            filtered_df = filtered_df.sort_values('period_duration', ascending=True)
        elif sort_by == "عدد اللوحات (الأكثر)":
            filtered_df = filtered_df.sort_values('total_boards', ascending=False)
        elif sort_by == "عدد اللوحات (الأقل)":
            filtered_df = filtered_df.sort_values('total_boards', ascending=True)
        else:
            filtered_df = filtered_df.sort_values('company_name')
        
        st.markdown(f"**📊 عرض {len(filtered_df)} شركة من أصل {total_companies}**")
        st.markdown("---")
        
        COLS_PER_ROW = 3
        companies_list = filtered_df.to_dict('records')
        
        for row_idx in range(0, len(companies_list), COLS_PER_ROW):
            row_companies = companies_list[row_idx:row_idx + COLS_PER_ROW]
            cols = st.columns(len(row_companies))
            
            for col_idx, company in enumerate(row_companies):
                with cols[col_idx]:
                    company_name = company['company_name']
                    total_boards = company['total_boards']
                    total_periods = company['total_periods']
                    period_duration = company.get('period_duration', 0)
                    first_period = company.get('first_period', '')
                    last_period = company.get('last_period', '')
                    year = company.get('العام', '')
                    
                    rank = row_idx + col_idx + 1
                    
                    if period_duration >= 10:
                        card_color = "linear-gradient(135deg, #667eea, #764ba2)"
                        status_icon = "🌟"
                        status_text = "مميز"
                    elif period_duration >= 5:
                        card_color = "linear-gradient(135deg, #11998e, #38ef7d)"
                        status_icon = "⭐"
                        status_text = "نشط"
                    else:
                        card_color = "linear-gradient(135deg, #f093fb, #f5576c)"
                        status_icon = "🔄"
                        status_text = "قريب"
                    
                    st.markdown(f'''
                    <div style="
                        background: {card_color};
                        border-radius: 20px;
                        padding: 20px;
                        margin-bottom: 20px;
                        margin-right: 16px;
                        margin-left: 16px;
                        color: white;
                        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
                        transition: all 0.3s ease;
                        height: 280px;
                        display: flex;
                        flex-direction: column;
                        justify-content: space-between;
                        cursor: pointer;
                    "
                    onmouseover="this.style.transform='translateY(-5px) scale(1.02)'"
                    onmouseout="this.style.transform='translateY(0) scale(1)'"
                    >
                        <div>
                            <div style="display: flex; justify-content: space-between; align-items: start;">
                                <h3 style="margin: 0; font-size: 18px; font-weight: bold;">
                                    #{rank} {status_icon} {company_name}
                                </h3>
                                <span style="background: rgba(255,255,255,0.2); padding: 4px 12px; border-radius: 20px; font-size: 12px;">{year}</span>
                            </div>
                            <div style="margin-top: 10px; font-size: 13px; opacity: 0.9;">
                                📅 {first_period} → {last_period}
                            </div>
                            <div style="margin-top: 4px; font-size: 12px; opacity: 0.8;">
                                ⏱️ المدة: {period_duration} فترة ({period_duration/2:.1f} شهر) | {status_text}
                            </div>
                        </div>
                        <div style="display: flex; justify-content: space-around; padding: 10px 0;">
                            <div style="text-align: center;">
                                <div style="font-size: 28px; font-weight: bold;">{total_boards}</div>
                                <div style="font-size: 12px; opacity: 0.8;">📊 لوحات</div>
                            </div>
                            <div style="text-align: center;">
                                <div style="font-size: 28px; font-weight: bold;">{total_periods}</div>
                                <div style="font-size: 12px; opacity: 0.8;">📅 فترات</div>
                            </div>
                            <div style="text-align: center;">
                                <div style="font-size: 28px; font-weight: bold;">{period_duration}</div>
                                <div style="font-size: 12px; opacity: 0.8;">⏱️ نصف شهر</div>
                            </div>
                        </div>
                        <div style="display: flex; gap: 8px; flex-wrap: wrap;">
                            <span style="background: rgba(255,255,255,0.2); padding: 2px 10px; border-radius: 12px; font-size: 11px;">🟢 نشط</span>
                            <span style="background: rgba(255,255,255,0.2); padding: 2px 10px; border-radius: 12px; font-size: 11px;">📌 {total_boards} لوحة</span>
                        </div>
                    </div>
                    ''', unsafe_allow_html=True)
                    
                    col_btn1, col_btn2 = st.columns(2)
                    with col_btn1:
                        if st.button(f"📋 تفاصيل", key=f"details_{company_name}_{row_idx}_{col_idx}", use_container_width=True):
                            st.session_state['selected_company'] = company_name
                            st.session_state['show_company_details'] = True
                            safe_rerun()
                    with col_btn2:
                        if st.button(f"🗺️ خريطة", key=f"map_{company_name}_{row_idx}_{col_idx}", use_container_width=True):
                            st.session_state['selected_company'] = company_name
                            st.session_state['show_company_map'] = True
                            safe_rerun()
    else:
        st.warning("⚠️ لا توجد شركات معلنة حالياً")
        st.info("💡 ستظهر الشركات هنا عندما يكون لديها حجوزات في الفترات الحالية أو المستقبلية")
    
    # ============================================================
    # باقي الكود (تفاصيل الشركة والخريطة) - كما هو
    # ============================================================
    
    if st.session_state.get('show_company_details', False):
        company_name = st.session_state.get('selected_company', '')
        if company_name:
            st.divider()
            st.subheader(f"📋 تفاصيل شركة {company_name}")
            
            @st.cache_data(ttl=120)
            def get_company_details(name):
                try:
                    conn = get_connection()
                    bookings_detail = pd.read_sql_query('''
                        SELECT "رقم اللوحة", "فترة الحجز", "العام"
                        FROM "حجوزات1" 
                        WHERE "اسم الزبون" = %s
                        ORDER BY "العام" DESC, "فترة الحجز"
                    ''', conn, params=(name,))
                    
                    if not bookings_detail.empty:
                        board_numbers = bookings_detail['رقم اللوحة'].unique().tolist()
                        placeholders = ','.join([f"'{b}'" for b in board_numbers])
                        boards_info = pd.read_sql_query(f'''
                            SELECT "رقم اللوحة", "اسم العمود" as location, "المحافظة" as city, "الشبكة" as network, "الحجم" as size
                            FROM "اعمدة انارة"
                            WHERE "رقم اللوحة" IN ({placeholders})
                        ''', conn)
                    else:
                        boards_info = pd.DataFrame()
                    
                    conn.close()
                    return bookings_detail, boards_info
                except Exception as e:
                    st.error(f"❌ خطأ: {str(e)}")
                    return pd.DataFrame(), pd.DataFrame()
            
            with st.spinner("🔄 جاري تحميل التفاصيل..."):
                bookings_detail, boards_info = get_company_details(company_name)
            
            if not bookings_detail.empty:
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("📊 إجمالي اللوحات", bookings_detail['رقم اللوحة'].nunique())
                with col2:
                    st.metric("📅 إجمالي الفترات", bookings_detail['فترة الحجز'].nunique())
                with col3:
                    st.metric("📆 السنوات", bookings_detail['العام'].nunique())
                
                st.write("**📋 تفاصيل الحجوزات:**")
                if not boards_info.empty:
                    merged = bookings_detail.merge(boards_info, on='رقم اللوحة', how='left')
                    st.dataframe(
                        merged[['رقم اللوحة', 'location', 'city', 'network', 'size', 'فترة الحجز', 'العام']],
                        use_container_width=True,
                        height=300
                    )
                else:
                    st.dataframe(bookings_detail, use_container_width=True, height=300)
                
                if st.button("🔙 إغلاق التفاصيل", use_container_width=True):
                    st.session_state['show_company_details'] = False
                    safe_rerun()
            else:
                st.info("📭 لا توجد تفاصيل متاحة")
    
    if st.session_state.get('show_company_map', False):
        company_name = st.session_state.get('selected_company', '')
        if company_name:
            st.divider()
            st.subheader(f"🗺️ خريطة مواقع شركة {company_name}")
            
            @st.cache_data(ttl=120)
            def get_company_locations(name):
                try:
                    conn = get_connection()
                    locations = pd.read_sql_query('''
                        SELECT DISTINCT
                            b."رقم اللوحة",
                            c."اسم العمود",
                            c."المحافظة",
                            c."الشبكة",
                            c."الحجم",
                            c."Latitude",
                            c."Longitude"
                        FROM "حجوزات1" b
                        JOIN "اعمدة انارة" c ON b."رقم اللوحة" = c."رقم اللوحة"
                        WHERE b."اسم الزبون" = %s
                        AND c."Latitude" IS NOT NULL
                        AND c."Longitude" IS NOT NULL
                    ''', conn, params=(name,))
                    conn.close()
                    return locations
                except Exception as e:
                    st.error(f"❌ خطأ: {str(e)}")
                    return pd.DataFrame()
            
            with st.spinner("🔄 جاري تحميل الخريطة..."):
                locations = get_company_locations(company_name)
            
            if not locations.empty:
                locations['Latitude'] = pd.to_numeric(locations['Latitude'], errors='coerce')
                locations['Longitude'] = pd.to_numeric(locations['Longitude'], errors='coerce')
                
                valid_locations = locations[
                    (locations['Latitude'].notna()) & 
                    (locations['Latitude'] != 0) &
                    (locations['Longitude'].notna()) & 
                    (locations['Longitude'] != 0)
                ]
                
                if not valid_locations.empty:
                    import folium
                    from streamlit_folium import st_folium
                    
                    center_lat = valid_locations['Latitude'].mean()
                    center_lon = valid_locations['Longitude'].mean()
                    
                    m = folium.Map(
                        location=[center_lat, center_lon],
                        zoom_start=8,
                        tiles='OpenStreetMap'
                    )
                    
                    for _, row in valid_locations.iterrows():
                        popup_text = f"""
                        <div dir="rtl" style="text-align:right; min-width:200px; font-family: Arial;">
                            <b>📍 {row['اسم العمود']}</b><br>
                            🏙️ {row['المحافظة']}<br>
                            📏 {row['الحجم']}<br>
                            📡 {row['الشبكة']}<br>
                            🆔 {row['رقم اللوحة']}
                        </div>
                        """
                        
                        folium.Marker(
                            location=[row['Latitude'], row['Longitude']],
                            popup=folium.Popup(popup_text, max_width=300),
                            tooltip=row['اسم العمود'],
                            icon=folium.Icon(color='green', icon='info-sign')
                        ).add_to(m)
                    
                    st_folium(m, width="100%", height=500)
                    
                    with st.expander("📍 قائمة المواقع", expanded=False):
                        st.dataframe(
                            valid_locations[['رقم اللوحة', 'اسم العمود', 'المحافظة', 'الشبكة', 'الحجم']],
                            use_container_width=True
                        )
                else:
                    st.info("📍 لا توجد إحداثيات صالحة لعرضها على الخريطة")
            else:
                st.warning("⚠️ لا توجد مواقع لهذه الشركة")
            
            if st.button("🔙 إغلاق الخريطة", use_container_width=True):
                st.session_state['show_company_map'] = False
                safe_rerun()
            




# ============================================================
# صفحة: الأعمدة المتاحة 
# ============================================================

elif page == "📍 الأعمدة المتاحة":
    st.title("📍 الأعمدة المتاحة للإيجار")
    st.info("📌 عرض الأعمدة حسب حالة الإتاحة مع عدد اللوحات الفعلية")
    
    # فلتر تاريخ البداية (بدون إعادة تحميل تلقائي)
    with st.form(key="filter_form"):
        st.subheader("📅 فلتر تاريخ بداية الإتاحة")
        start_date = st.date_input(
            "عرض الأعمدة المتاحة من تاريخ:",
            value=date.today(),
            help="اختر التاريخ الذي تبدأ منه فترة الإتاحة"
        )
        submitted = st.form_submit_button("🔍 تطبيق الفلتر")
    
    if not submitted and 'df' not in st.session_state:
        submitted = True
    
    if submitted:
        # حساب الفترة المستهدفة
        target_period_num = get_period_from_date(start_date)
        target_year = start_date.year
        st.write(f"📅 التاريخ المختار: {start_date}")
        st.write(f"📅 رقم الفترة: {target_period_num}")
        # جلب البيانات مع تخزين مؤقت
        @st.cache_data(ttl=300)
        def load_data(target_period_num, target_year):
            conn = get_connection()
            
            query = f"""
            WITH booking_periods AS (
                SELECT 
                    CAST("رقم اللوحة" AS TEXT) as "رقم اللوحة",
                    "فترة الحجز",
                    "العام",
                    CASE
                        WHEN "فترة الحجز" = 'كانون ثاني 15-1' THEN 1
                        WHEN "فترة الحجز" = 'كانون ثاني 30-15' THEN 2
                        WHEN "فترة الحجز" = 'شباط 15-1' THEN 3
                        WHEN "فترة الحجز" = 'شباط 30-15' THEN 4
                        WHEN "فترة الحجز" = 'اذار 15-1' THEN 5
                        WHEN "فترة الحجز" = 'اذار 30-15' THEN 6
                        WHEN "فترة الحجز" = 'نيسان 15-1' THEN 7
                        WHEN "فترة الحجز" = 'نيسан 30-15' THEN 8
                        WHEN "فترة الحجز" = 'ايار15-1' THEN 9
                        WHEN "فترة الحجز" = 'أيار 30-15' THEN 10
                        WHEN "فترة الحجز" = 'حزيران 15-1' THEN 11
                        WHEN "فترة الحجز" = 'حزيران 30-15' THEN 12
                        WHEN "فترة الحجز" = 'تموز 15-1' THEN 13
                        WHEN "فترة الحجز" = 'تموز 30-15' THEN 14
                        WHEN "فترة الحجز" = 'اب 15-1' THEN 15
                        WHEN "فترة الحجز" = 'اب 30-15' THEN 16
                        WHEN "فترة الحجز" = 'أيلول 15-1' THEN 17
                        WHEN "فترة الحجز" = 'ايلول30-15' THEN 18
                        WHEN "فترة الحجز" = 'تشرين اول 15-1' THEN 19
                        WHEN "فترة الحجز" = 'تشرين اول30-15' THEN 20
                        WHEN "فترة الحجز" = 'تشرين ثاني 15-1' THEN 21
                        WHEN "فترة الحجز" = 'تشرين ثاني 30-15' THEN 22
                        WHEN "فترة الحجز" = 'كانون اول 15-1' THEN 23
                        WHEN "فترة الحجز" = 'كانون اول 30-15' THEN 24
                    END as period_num
                FROM "حجوزات1"
                WHERE "العام" >= {target_year}
            ),
            board_aggregated AS (
                SELECT 
                    "رقم اللوحة",
                    MAX(CASE WHEN "العام" = {target_year} AND "period_num" = {target_period_num} THEN 1 ELSE 0 END) as has_current,
                    MAX(CASE WHEN ("العام" > {target_year}) OR ("العام" = {target_year} AND "period_num" > {target_period_num}) THEN 1 ELSE 0 END) as has_future,
                    MIN(CASE WHEN ("العام" > {target_year}) OR ("العام" = {target_year} AND "period_num" > {target_period_num}) THEN period_num ELSE NULL END) as min_future_period,
                    MAX(CASE WHEN "period_num" <= {target_period_num} THEN period_num ELSE NULL END) as max_current_period
                FROM booking_periods
                GROUP BY "رقم اللوحة"
            )
            SELECT 
                a."رقم اللوحة",
                a."اسم العمود",
                a."المحافظة",
                a."الشبكة",
                a."الحجم",
                a."العدد",
                CASE 
                    WHEN b.has_current = 1 AND b.has_future = 1 THEN '🔴 محجوز بالكامل'
                    WHEN b.has_current = 1 AND b.has_future = 0 THEN '🟠 محجوز مؤقتاً'
                    WHEN b.has_current = 0 AND b.has_future = 1 THEN '🟡 متاح مؤقتاً'
                    ELSE '🟢 متاح فوراً'
                END as status,
                b.min_future_period as next_booking_period,
                b.max_current_period as end_booking_period
            FROM "اعمدة انارة" a
            LEFT JOIN board_aggregated b ON CAST(a."رقم اللوحة" AS TEXT) = b."رقم اللوحة"
            ORDER BY a."المحافظة", a."رقم اللوحة"
            """
            
            df = pd.read_sql_query(query, conn)
            conn.close()
            return df
        
        df = load_data(target_period_num, target_year)
        
        # حساب الإحصائيات
        available_now_sites = len(df[df['status'] == '🟢 متاح فوراً'])
        available_now_boards = df[df['status'] == '🟢 متاح فوراً']['العدد'].sum()
        
        available_temp_sites = len(df[df['status'] == '🟡 متاح مؤقتاً'])
        available_temp_boards = df[df['status'] == '🟡 متاح مؤقتاً']['العدد'].sum()
        
        booked_temp_sites = len(df[df['status'] == '🟠 محجوز مؤقتاً'])
        booked_temp_boards = df[df['status'] == '🟠 محجوز مؤقتاً']['العدد'].sum()
        
        booked_full_sites = len(df[df['status'] == '🔴 محجوز بالكامل'])
        booked_full_books = df[df['status'] == '🔴 محجوز بالكامل']['العدد'].sum()
        
        # ✅ المجاميع الكلية
        total_available_sites = available_now_sites + available_temp_sites
        total_available_boards = available_now_boards + available_temp_boards
        
        total_booked_sites = booked_temp_sites + booked_full_sites
        total_booked_boards = booked_temp_boards + booked_full_books
        
        total_sites = len(df)
        total_boards = df['العدد'].sum()
        
        # عرض الإحصائيات
        st.subheader("📊 إحصائيات عامة")
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("#### 🟢 متاح فوراً")
            st.markdown(f"📍 **المواقع:** {available_now_sites}")
            st.markdown(f"📌 **اللوحات:** {int(available_now_boards):,}")
        
        with col2:
            st.markdown("#### 🟡 متاح مؤقتاً")
            st.markdown(f"📍 **المواقع:** {available_temp_sites}")
            st.markdown(f"📌 **اللوحات:** {int(available_temp_boards):,}")
        
        col3, col4 = st.columns(2)
        with col3:
            st.markdown("#### 🟠 محجوز مؤقتاً")
            st.markdown(f"📍 **المواقع:** {booked_temp_sites}")
            st.markdown(f"📌 **اللوحات:** {int(booked_temp_boards):,}")
        
        with col4:
            st.markdown("#### 🔴 محجوز بالكامل")
            st.markdown(f"📍 **المواقع:** {booked_full_sites}")
            st.markdown(f"📌 **اللوحات:** {int(booked_full_books):,}")
        
        st.divider()
        
        # ✅ عرض المجاميع الكلية
        col_total1, col_total2, col_total3 = st.columns(3)
        with col_total1:
            st.markdown(f"#### ✅ إجمالي المتاح")
            st.markdown(f"📍 **المواقع:** {total_available_sites}")
            st.markdown(f"📌 **اللوحات:** {int(total_available_boards):,}")
        
        with col_total2:
            st.markdown(f"#### 🔴 إجمالي المحجوز")
            st.markdown(f"📍 **المواقع:** {total_booked_sites}")
            st.markdown(f"📌 **اللوحات:** {int(total_booked_boards):,}")
        
        with col_total3:
            st.markdown(f"#### 📊 الإجمالي الكلي")
            st.markdown(f"📍 **المواقع:** {total_sites}")
            st.markdown(f"📌 **اللوحات:** {int(total_boards):,}")
        
        st.divider()
        
        # عرض حسب المحافظة
        for city in df['المحافظة'].unique():
            city_data = df[df['المحافظة'] == city]
            
            with st.expander(f"🏙️ {city} - {len(city_data)} موقع", expanded=False):  # expanded=False يقلل التحميل
                
                display_df = city_data.copy()
                
                def period_to_text(period_num):
                    if pd.isna(period_num):
                        return ""
                    period_map = {11: "يبدأ 1/6", 12: "يبدأ 16/6", 13: "يبدأ 1/7", 14: "يبدأ 16/7"}
                    return period_map.get(period_num, f"فترة {period_num}")
                
                display_df['تاريخ البدء'] = display_df['next_booking_period'].apply(period_to_text)
                display_df['تاريخ الانتهاء'] = display_df['end_booking_period'].apply(period_to_text)
                
                # استخدام use_container_width=False لتقليل إعادة التحميل
                st.dataframe(
                    display_df[['رقم اللوحة', 'اسم العمود', 'الشبكة', 'الحجم', 'العدد', 'status', 'تاريخ البدء', 'تاريخ الانتهاء']],
                    use_container_width=False,
                    height=300
                )
        
        # تصدير
        csv_data = df[['رقم اللوحة', 'اسم العمود', 'المحافظة', 'الشبكة', 'الحجم', 'العدد', 'status', 'next_booking_period', 'end_booking_period']].to_csv(
            index=False, encoding='utf-8-sig'
        )
        st.download_button(
            "📥 تحميل التقرير (CSV)",
            csv_data,
            f"available_boards_{start_date.strftime('%Y%m%d')}.csv",
            "text/csv",
            use_container_width=True
        )

# ============================================================
# صفحة: لوحة التحكم البصرية للفترات
# ============================================================

elif page == "📅 لوحة الفترات":
    st.title("📅 لوحة التحكم البصرية للفترات")
    st.info("📌 عرض المتاح والمحجوز لكل فترة مع تفاصيل اللوحات المتاحة")
    
    # ============================================================
    # بناء PERIOD_ORDER من قاعدة البيانات
    # ============================================================
    
    def build_period_order():
        conn = get_connection()
        df = pd.read_sql_query('SELECT no, namee FROM "الفترة" ORDER BY no', conn)
        conn.close()
        return {row['namee']: row['no'] for _, row in df.iterrows()}
    
    PERIOD_ORDER = build_period_order()
    
    # ============================================================
    # الفلاتر
    # ============================================================
    
    col_filter1, col_filter2 = st.columns(2)
    
    with col_filter1:
        conn = get_connection()
        cities_df = pd.read_sql_query('SELECT DISTINCT "المحافظة" FROM "اعمدة انارة" ORDER BY "المحافظة"', conn)
        conn.close()
        city_list = ['جميع المحافظات'] + cities_df['المحافظة'].tolist()
        selected_city = st.selectbox("🏙️ اختر المحافظة:", city_list)
    
    with col_filter2:
        conn = get_connection()
        sizes_df = pd.read_sql_query('SELECT DISTINCT "الحجم" FROM "اعمدة انارة" ORDER BY "الحجم"', conn)
        conn.close()
        size_list = ['جميع الأحجام'] + sizes_df['الحجم'].tolist()
        selected_size = st.selectbox("📏 اختر الحجم:", size_list)
    
    # ============================================================
    # جلب البيانات
    # ============================================================
    
    @st.cache_data(ttl=300)
    def load_period_data(selected_city, selected_size, PERIOD_ORDER):
        conn = get_connection()
        
        where_conditions = []
        if selected_city != 'جميع المحافظات':
            where_conditions.append(f'"المحافظة" = \'{selected_city}\'')
        if selected_size != 'جميع الأحجام':
            where_conditions.append(f'"الحجم" = \'{selected_size}\'')
        
        where_clause = " AND ".join(where_conditions) if where_conditions else "1=1"
        
        boards_query = f"""
        SELECT 
            CAST("رقم اللوحة" AS TEXT) as "رقم اللوحة",
            "اسم العمود",
            "المحافظة",
            "الشبكة",
            "الحجم",
            "العدد"
        FROM "اعمدة انارة"
        WHERE {where_clause}
        """
        boards_df = pd.read_sql_query(boards_query, conn)
        
        bookings_query = """
        SELECT 
            CAST("رقم اللوحة" AS TEXT) as "رقم اللوحة",
            "اسم الزبون",
            "فترة الحجز",
            "العام"
        FROM "حجوزات1"
        WHERE "العام" = 2026
        """
        bookings_df = pd.read_sql_query(bookings_query, conn)
        conn.close()
        
        bookings_df['period_num'] = bookings_df['فترة الحجز'].apply(
            lambda x: PERIOD_ORDER.get(x, 99)
        )
        
        return boards_df, bookings_df
    
    boards_df, bookings_df = load_period_data(selected_city, selected_size, PERIOD_ORDER)
    
    # ============================================================
    # الحصول على الفترات من PERIOD_ORDER
    # ============================================================
    
    sorted_periods = sorted(PERIOD_ORDER.items(), key=lambda x: x[1])
    all_period_names = [p[0] for p in sorted_periods]
    
    # ============================================================
    # حساب الإحصائيات لكل فترة
    # ============================================================
    
    total_boards = boards_df['العدد'].sum()
    period_stats = []
    period_details = {}
    
    for period_name, period_num in sorted_periods:
        # اللوحات المحجوزة في هذه الفترة
        booked_boards = bookings_df[bookings_df['period_num'] == period_num]['رقم اللوحة'].unique()
        booked_boards_list = list(booked_boards)
        
        # تفاصيل اللوحات المحجوزة
        booked_details = boards_df[boards_df['رقم اللوحة'].isin(booked_boards_list)]
        
        # ✅ اللوحات المتاحة في هذه الفترة (جميع اللوحات - المحجوزة)
        available_boards = boards_df[~boards_df['رقم اللوحة'].isin(booked_boards_list)]
        
        # الزبائن في هذه الفترة
        customers = bookings_df[bookings_df['period_num'] == period_num]['اسم الزبون'].unique()
        customers_list = ', '.join(customers[:3]) + (f' و {len(customers)-3} آخرين' if len(customers) > 3 else '')
        
        period_stats.append({
            'الفترة': period_name,
            'رقم الفترة': period_num,
            'إجمالي اللوحات': int(total_boards),
            'محجوز': int(booked_details['العدد'].sum()),
            'متاح': int(total_boards - booked_details['العدد'].sum()),
            'عدد الزبائن': len(customers),
            'الزبائن': customers_list if len(customers) > 0 else 'لا يوجد'
        })
        
        period_details[period_name] = {
            'booked_details': booked_details,
            'available_details': available_boards,  # ✅ اللوحات المتاحة
            'customers': customers,
            'booked_boards': booked_boards_list
        }
    
    period_df = pd.DataFrame(period_stats)
    
    # ============================================================
    # عرض النتائج
    # ============================================================
    
    st.subheader("📊 إحصائيات عامة")
    col1, col2, col3 = st.columns(3)
    col1.metric("🏢 إجمالي اللوحات", int(total_boards))
    col2.metric("📅 عدد الفترات", len(all_period_names))
    col3.metric("👥 عدد الزبائن", bookings_df['اسم الزبون'].nunique())
    
    st.divider()
    
    st.subheader("📋 الفترات")
    
    for i in range(0, len(all_period_names), 4):
        cols = st.columns(4)
        for j, col in enumerate(cols):
            if i + j < len(all_period_names):
                period_name = all_period_names[i + j]
                stats = period_stats[i + j]
                
                with col:
                    if stats['محجوز'] == 0:
                        bg_color = "#e8f5e9"
                        border_color = "#4CAF50"
                    else:
                        bg_color = "#fff3e0"
                        border_color = "#FF9800"
                    
                    st.markdown(f"""
                    <div style="background:{bg_color};border:2px solid {border_color};border-radius:12px;padding:15px;text-align:center;margin:5px 0;">
                        <div style="font-size:14px;font-weight:bold;">{period_name}</div>
                        <div style="font-size:12px;margin-top:5px;">🟢 {stats['متاح']} | 🔴 {stats['محجوز']}</div>
                        <div style="font-size:11px;margin-top:5px;color:#666;">👥 {stats['الزبائن']}</div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if st.button(f"📋 تفاصيل", key=f"detail_{i+j}"):
                        st.session_state[f'selected_period_{i+j}'] = period_name
                        st.session_state['show_period_detail'] = True
                        st.rerun()
    
    # ============================================================
    # التفاصيل للفترة المختارة (عرض المتاح)
    # ============================================================
    
    if st.session_state.get('show_period_detail', False):
        selected_period = None
        for key in st.session_state:
            if key.startswith('selected_period_'):
                selected_period = st.session_state[key]
                break
        
        if selected_period and selected_period in period_details:
            details = period_details[selected_period]
            
            st.divider()
            st.subheader(f"📋 تفاصيل الفترة: {selected_period}")
            
            period_stat = next(p for p in period_stats if p['الفترة'] == selected_period)
            col1, col2, col3 = st.columns(3)
            col1.metric("📊 إجمالي اللوحات", period_stat['إجمالي اللوحات'])
            col2.metric("🔴 محجوز", period_stat['محجوز'])
            col3.metric("🟢 متاح", period_stat['متاح'])
            
            # عرض الزبائن في هذه الفترة
            if len(details['customers']) > 0:
                st.write("**👥 الزبائن في هذه الفترة:**")
                st.write(", ".join(details['customers']))
            else:
                st.write("**👥 الزبائن في هذه الفترة:** لا يوجد")
            
            # ✅ عرض اللوحات المتاحة (وليس المحجوزة)
            if not details['available_details'].empty:
                st.write("**📋 اللوحات المتاحة في هذه الفترة:**")
                st.dataframe(
                    details['available_details'][['رقم اللوحة', 'اسم العمود', 'المحافظة', 'الشبكة', 'الحجم', 'العدد']],
                    use_container_width=True
                )
            else:
                st.info("✅ لا توجد لوحات متاحة في هذه الفترة")
            
            if st.button("🔙 إغلاق التفاصيل", key="close_detail"):
                st.session_state['show_period_detail'] = False
                for key in list(st.session_state.keys()):
                    if key.startswith('selected_period_'):
                        del st.session_state[key]
                st.rerun()
    
    # ============================================================
    # الرسم البياني
    # ============================================================
    
    st.divider()
    st.subheader("📊 رسم بياني للمتاح والمحجوز")
    
    fig = go.Figure()
    fig.add_trace(go.Bar(x=period_df['الفترة'], y=period_df['متاح'], name='متاح', marker_color='#4CAF50'))
    fig.add_trace(go.Bar(x=period_df['الفترة'], y=period_df['محجوز'], name='محجوز', marker_color='#f44336'))
    fig.update_layout(
        barmode='stack',
        height=400,
        xaxis_tickangle=-45,
        xaxis_title='الفترة',
        yaxis_title='عدد اللوحات'
    )
    st.plotly_chart(fig, use_container_width=True)
    
    # ============================================================
    # تصدير Excel (مع المتاح في التفاصيل)
    # ============================================================
    
    st.divider()
    st.subheader("📥 تصدير التقرير")
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # صفحة الملخص
        period_df.to_excel(writer, sheet_name='ملخص الفترات', index=False)
        
        # صفحة لكل فترة تحتوي على اللوحات المتاحة
        for period_name in all_period_names:
            if period_name in period_details:
                details = period_details[period_name]
                if not details['available_details'].empty:
                    sheet_name = period_name[:25]
                    details['available_details'].to_excel(writer, sheet_name=sheet_name, index=False)
    
    output.seek(0)
    
    st.download_button(
        "📥 تحميل تقرير Excel",
        output,
        f"periods_report_{selected_city}_{selected_size}.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )
#=================
# لوحة المراقبة
#==================


elif page == "📊 Dashboard":
    st.markdown("""
    <div style="text-align: center; margin-bottom: 30px;">
        <h1>📊 لوحة التحكم المتقدمة</h1>
        <p style="color: rgba(255,255,255,0.7);">نظرة شاملة على أداء النظام وإحصائيات الإعلانات</p>
    </div>
    """, unsafe_allow_html=True)
    
    current_year = datetime.now().year
    today = date.today()
    target_period_num = get_period_from_date(today)
    
    # ============================================================
    # جلب البيانات (نفس منطق صفحة الأعمدة المتاحة)
    # ============================================================
    
    @st.cache_data(ttl=300)
    def load_dashboard_data(target_period_num, target_year):
        conn = get_connection()
        
        query = f"""
        WITH booking_periods AS (
            SELECT 
                CAST("رقم اللوحة" AS TEXT) as "رقم اللوحة",
                "فترة الحجز",
                "العام",
                CASE
                    WHEN "فترة الحجز" = 'كانون ثاني 15-1' THEN 1
                    WHEN "فترة الحجز" = 'كانون ثاني 30-15' THEN 2
                    WHEN "فترة الحجز" = 'شباط 15-1' THEN 3
                    WHEN "فترة الحجز" = 'شباط 30-15' THEN 4
                    WHEN "فترة الحجز" = 'اذار 15-1' THEN 5
                    WHEN "فترة الحجز" = 'اذار 30-15' THEN 6
                    WHEN "فترة الحجز" = 'نيسان 15-1' THEN 7
                    WHEN "فترة الحجز" = 'نيسان 30-15' THEN 8
                    WHEN "فترة الحجز" = 'ايار15-1' THEN 9
                    WHEN "فترة الحجز" = 'أيار 30-15' THEN 10
                    WHEN "فترة الحجز" = 'حزيران 15-1' THEN 11
                    WHEN "فترة الحجز" = 'حزيران 30-15' THEN 12
                    WHEN "فترة الحجز" = 'تموز 15-1' THEN 13
                    WHEN "فترة الحجز" = 'تموز 30-15' THEN 14
                    WHEN "فترة الحجز" = 'اب 15-1' THEN 15
                    WHEN "فترة الحجز" = 'اب 30-15' THEN 16
                    WHEN "فترة الحجز" = 'أيلول 15-1' THEN 17
                    WHEN "فترة الحجز" = 'ايلول30-15' THEN 18
                    WHEN "فترة الحجز" = 'تشرين اول 15-1' THEN 19
                    WHEN "فترة الحجز" = 'تشرين اول30-15' THEN 20
                    WHEN "فترة الحجز" = 'تشرين ثاني 15-1' THEN 21
                    WHEN "فترة الحجز" = 'تشرين ثاني 30-15' THEN 22
                    WHEN "فترة الحجز" = 'كانون اول 15-1' THEN 23
                    WHEN "فترة الحجز" = 'كانون اول 30-15' THEN 24
                END as period_num
            FROM "حجوزات1"
            WHERE "العام" >= {target_year}
        ),
        board_aggregated AS (
            SELECT 
                "رقم اللوحة",
                MAX(CASE WHEN "العام" = {target_year} AND "period_num" = {target_period_num} THEN 1 ELSE 0 END) as has_current,
                MAX(CASE WHEN ("العام" > {target_year}) OR ("العام" = {target_year} AND "period_num" > {target_period_num}) THEN 1 ELSE 0 END) as has_future
            FROM booking_periods
            GROUP BY "رقم اللوحة"
        )
        SELECT 
            a."رقم اللوحة",
            a."اسم العمود",
            a."المحافظة",
            a."الشبكة",
            a."الحجم",
            a."العدد",
            CASE 
                WHEN b.has_current = 1 AND b.has_future = 1 THEN 'محجوز بالكامل'
                WHEN b.has_current = 1 AND b.has_future = 0 THEN 'محجوز مؤقتاً'
                WHEN b.has_current = 0 AND b.has_future = 1 THEN 'متاح مؤقتاً'
                ELSE 'متاح فوراً'
            END as status
        FROM "اعمدة انارة" a
        LEFT JOIN board_aggregated b ON CAST(a."رقم اللوحة" AS TEXT) = b."رقم اللوحة"
        ORDER BY a."المحافظة", a."رقم اللوحة"
        """
        
        df = pd.read_sql_query(query, conn)
        conn.close()
        return df
    
    df = load_dashboard_data(target_period_num, current_year)
    
    # ============================================================
    # حساب الإحصائيات (نفس صفحة الأعمدة المتاحة)
    # ============================================================
    
    total_sites = len(df)
    total_boards = df['العدد'].sum()
    
    # محجوز حالياً (has_current = 1)
    booked_current = df[df['status'].isin(['محجوز مؤقتاً', 'محجوز بالكامل'])]
    booked_current_sites = len(booked_current)
    booked_current_boards = booked_current['العدد'].sum()
    
    # متاح حالياً (has_current = 0)
    available_current = df[~df['status'].isin(['محجوز مؤقتاً', 'محجوز بالكامل'])]
    available_current_sites = len(available_current)
    available_current_boards = available_current['العدد'].sum()
    
    # نسبة الإشغال
    occupancy_rate = (booked_current_sites / total_sites * 100) if total_sites > 0 else 0
    
    # ============================================================
    # عرض البطاقات
    # ============================================================
    
    cols = st.columns(4)
    metrics_data = [
        ("إجمالي المواقع", total_sites, "🗺️", "primary"),
        ("🔴 محجوز", booked_current_sites, "📌", "danger"),
        ("🟢 متاح", available_current_sites, "✅", "success"),
        ("📈 نسبة الإشغال", f"{occupancy_rate:.1f}%", "📊", "warning")
    ]
    
    for idx, (title, value, icon, color) in enumerate(metrics_data):
        with cols[idx]:
            st.markdown(create_metric_card_3d(title, value, icon, color), unsafe_allow_html=True)
    
    # ============================================================
    # شريط التقدم
    # ============================================================
    
    st.markdown(f"""
    <div style="margin: 20px 0;">
        <div style="display: flex; justify-content: space-between; margin-bottom: 8px;">
            <span>📊 نسبة الإشغال الحالية</span>
            <span style="font-weight: bold;">{occupancy_rate:.1f}%</span>
        </div>
        <div style="height: 12px; background: rgba(0,0,0,0.1); border-radius: 10px; overflow: hidden;">
            <div style="width: {occupancy_rate}%; height: 100%; background: linear-gradient(90deg, #667eea, #764ba2); border-radius: 10px;"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.divider()
    
    # ============================================================
    # إحصائيات اللوحات الفعلية
    # ============================================================
    
    st.subheader("📊 إحصائيات اللوحات الفعلية")
    col_boards1, col_boards2, col_boards3 = st.columns(3)
    with col_boards1:
        st.metric("📌 إجمالي اللوحات", f"{int(total_boards):,}")
    with col_boards2:
        st.metric("🔴 لوحات محجوزة", f"{int(booked_current_boards):,}")
    with col_boards3:
        st.metric("🟢 لوحات متاحة", f"{int(available_current_boards):,}")
    
    st.divider()
    
    
    # ============================================================
    # الرسوم البيانية
    # ============================================================
    
    col_chart1, col_chart2 = st.columns(2)
    
    with col_chart1:
        st.subheader("🥧 نسبة الإشغال")
        fig_pie = go.Figure(data=[go.Pie(
            labels=['محجوز', 'متاح'],
            values=[booked_current_sites, available_current_sites],
            hole=0.4,
            marker_colors=['#dc2626', '#22c55e'],
            textinfo='percent+label'
        )])
        fig_pie.update_layout(height=400, paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)')
        st.plotly_chart(fig_pie, use_container_width=True)
    
    with col_chart2:
        st.subheader("📊 نسبة الإشغال حسب المحافظة")
        
        city_stats = []
        for city in df['المحافظة'].unique():
            city_data = df[df['المحافظة'] == city]
            city_total = len(city_data)
            city_booked = len(city_data[city_data['status'].isin(['محجوز مؤقتاً', 'محجوز بالكامل'])])
            city_stats.append({
                'المحافظة': city,
                'نسبة الإشغال': (city_booked / city_total * 100) if city_total > 0 else 0
            })
        
        city_df = pd.DataFrame(city_stats)
        
        # ✅ ألوان واضحة ومتباينة
        fig_bar = px.bar(
            city_df, 
            x='المحافظة', 
            y='نسبة الإشغال',
            color='نسبة الإشغال',
            color_continuous_scale=[
                (0.0, '#22c55e'),   # أخضر فاتح (نسبة منخفضة)
                (0.5, '#eab308'),   # أصفر (نسبة متوسطة)
                (1.0, '#dc2626')    # أحمر (نسبة عالية)
            ],
            text='نسبة الإشغال'
        )
        fig_bar.update_traces(
            texttemplate='%{text:.1f}%',
            textposition='outside',
            marker=dict(line=dict(width=2, color='#1e293b'))
        )
        fig_bar.update_layout(
            height=400,
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            font=dict(color='#1e293b', size=13),
            xaxis=dict(tickangle=0, tickfont=dict(size=11)),
            yaxis=dict(
                title='نسبة الإشغال %',
                range=[0, 105],
                gridcolor='rgba(0,0,0,0.08)'
            ),
            coloraxis_colorbar=dict(
                title='نسبة الإشغال',
                tickvals=[0, 50, 100],
                ticktext=['منخفضة', 'متوسطة', 'عالية']
            )
        )
        st.plotly_chart(fig_bar, use_container_width=True)
    
    st.divider()
    

    
    # ============================================================
    # الخريطة (مع منع الريفريش)
    # ============================================================
    
    # ============================================================
    # الخريطة (مع زر عرض/إخفاء)
    # ============================================================
    
    # ✅ استخدام st.expander لتوسيع الخريطة فقط عند الطلب
    with st.expander("🗺️ عرض الخريطة", expanded=False):
        
        @st.cache_data(ttl=600)
        def load_map_data():
            return run_query('SELECT * FROM "اعمدة انارة"')
        
        all_columns_map = load_map_data()
        
        # ✅ جلب أسماء الزبائن للوحات المحجوزة
        @st.cache_data(ttl=600)
        def load_customer_names():
            return run_query('SELECT DISTINCT "رقم اللوحة", "اسم الزبون" FROM "حجوزات1" WHERE "العام" = %s', (current_year,))
        
        customers_df = load_customer_names()
        customer_dict = dict(zip(customers_df['رقم اللوحة'].astype(str), customers_df['اسم الزبون']))
        
        # ✅ إنشاء الخريطة
        m = folium.Map(
            location=SYRIA_COORDS["سوريا"], 
            zoom_start=7,
            height=480
        )
        marker_cluster = MarkerCluster().add_to(m)
        
        booked_boards_list = df[df['status'].isin(['محجوز مؤقتاً', 'محجوز بالكامل'])]['رقم اللوحة'].tolist()
        
        for _, row in all_columns_map.iterrows():
            if pd.notnull(row.get('Latitude')) and pd.notnull(row.get('Longitude')) and row.get('Latitude') != 0:
                is_booked = row['رقم اللوحة'] in booked_boards_list
                board_id = str(row['رقم اللوحة'])
                
                if is_booked:
                    color = 'red'
                    status_text = '🔴 محجوز'
                    customer_name = customer_dict.get(board_id, '')
                    customer_line = f"👤 الزبون: {customer_name}<br>" if customer_name else ""
                else:
                    color = 'green'
                    status_text = '🟢 متاح'
                    customer_line = ""
                
                network = row.get('الشبكة', 0)
                network_display = str(int(network)) if network is not None and network != 0 else 'بدون شبكة'
                board_count = int(row['العدد'])
                
                popup_html = f"""
                <div dir="rtl" style="font-family:Arial;text-align:right;min-width:250px;padding:5px;">
                    <b>🏢 {row['اسم العمود']}</b><br>
                    📍 {row['المحافظة']}<br>
                    📡 الشبكة: {network_display}<br>
                    📏 {row['الحجم']}<br>
                    🔢 العدد: {board_count} لوحة<br>
                    {customer_line}
                    {status_text}
                </div>
                """
                
                folium.Marker(
                    [row['Latitude'], row['Longitude']],
                    popup=folium.Popup(popup_html, max_width=350),
                    icon=folium.Icon(color=color)
                ).add_to(marker_cluster)
        
        # ✅ عرض الخريطة داخل expander
        st_folium(m, width="100%", height=480, key="dashboard_map")
        
        # ✅ تذييل صغير
        st.caption("📍 اضغط على أيقونة لعرض تفاصيل اللوحة • 🟢 متاح • 🔴 محجوز")
#============================
# عرض سعر
#=============================

elif page == "📄 عرض سعر":
    st.title("📄 بناء عرض سعر جديد")
    st.markdown('<div class="custom-divider"></div>', unsafe_allow_html=True)
    
    try:
        with st.expander("🔔 العروض المنتهية (تحتاج إلى إجراء)", expanded=False):
            manage_expired_offers()
        
        st.subheader("📂 استرجاع عرض محفوظ")
        saved_offers = run_query('SELECT id, client_name, offer_date, start_p, end_p, year, status FROM "offers_history" WHERE status = %s ORDER BY id DESC', ('Pending',))
        
        if saved_offers is not None and not saved_offers.empty:
            offer_options = {}
            for _, row in saved_offers.iterrows():
                # طريقة آمنة تماماً للحصول على التاريخ
                offer_date = row['offer_date']
                try:
                    # محاولة التحويل إلى string
                    date_str = str(offer_date)[:10] if offer_date else "بدون تاريخ"
                except:
                    date_str = "بدون تاريخ"
                offer_options[f"{row['client_name']} ({date_str})"] = row['id']
            
            selected_offer = st.selectbox("اختر عرضاً محفوظاً:", ["---"] + list(offer_options.keys()), key="load_offer_select")
            
            if selected_offer != "---" and st.button("🔄 تحميل للسلة", key="load_offer_button", use_container_width=True):
                try:
                    offer_id = offer_options[selected_offer]
                    result = run_query('SELECT cart_json, client_name, start_p, end_p, year FROM "offers_history" WHERE id = %s', (offer_id,))
                    
                    if result is not None and not result.empty:
                        row = result.iloc[0]
                        data = json.loads(row['cart_json'])
                        cart_raw = data.get("data", data)
                        st.session_state.cart = {}
                        for city, networks in cart_raw.items():
                            st.session_state.cart[city] = {}
                            for net, df_dict in networks.items():
                                st.session_state.cart[city][net] = pd.DataFrame(df_dict)
                        
                        st.session_state.temp_cust = row['client_name']
                        st.success("✅ تم تحميل العرض بنجاح")
                        st.rerun()
                except Exception as e:
                    st.error(f"خطأ في تحميل العرض: {str(e)}")
        
        st.divider()
        
        draw_df = run_query('SELECT * FROM "اسماء الرسم"')
        
        customer_name = st.text_input("🏢 اسم الزبون", value=st.session_state.get('temp_cust', ""), placeholder="أدخل اسم الشركة أو الزبون")
        st.session_state.temp_cust = customer_name
        
        col1, col2, col3 = st.columns(3)
        with col1:
            selected_size = st.selectbox("📏 قياس اللوحة:", draw_df['الحجم'].unique().tolist())
        with col2:
            print_type = st.radio("🖨️ نوع الطباعة:", ["عادي", "سكوتش"], horizontal=True)
        with col3:
            year = st.number_input("📅 العام:", min_value=2024, max_value=2030, value=2026)
        
        is_foreign = st.checkbox("🌍 منتج أجنبي")
        
        periods_df = run_query('SELECT namee, no FROM "الفترة" ORDER BY no')
        period_names = periods_df['namee'].tolist()
        
        if not period_names:
            st.error("❌ لا توجد فترات في جدول الفترة")
            st.stop()
        
        col_p1, col_p2 = st.columns(2)
        with col_p1:
            start_p = st.selectbox("📅 من فترة:", period_names, key="start_period")
        with col_p2:
            end_p = st.selectbox("📅 إلى فترة:", period_names, index=len(period_names)-1, key="end_period")
        
        start_idx = period_names.index(start_p)
        end_idx = period_names.index(end_p)
        periods_count = abs(end_idx - start_idx) + 1
        months_count = periods_count / 2
        selected_periods = period_names[start_idx:end_idx+1]
        
        st.info(f"📅 عدد الفترات: {periods_count} | عدد الأشهر: {months_count:.1f}")
        
        fee_print, fee_ads = get_fees(draw_df, selected_size, print_type, is_foreign)
        
        per_column_print = fee_print
        per_column_display = fee_ads * months_count
        per_column_total = per_column_print + per_column_display
        
        st.success(f"""
        💰 **تفاصيل الأسعار:**
        - أجر الطباعة الثابت: **{fee_print}$**
        - أجر العرض الشهري: **{fee_ads}$**
        - المدة: **{months_count:.1f} شهر**
        - الإجمالي لكل عمود: **{per_column_total:.2f}$**
        """)
        
        st.divider()
        st.subheader("📍 اختيار المواقع")
        
        # ============================================================
        # 1. اختيار المحافظة والحجم
        # ============================================================
        
        cities = run_query('SELECT DISTINCT "المحافظة" FROM "اعمدة انارة"')['المحافظة'].tolist()
        selected_city = st.selectbox("اختر المحافظة:", cities)
        
        # ============================================================
        # 2. جلب جميع الأعمدة في المحافظة والحجم المختارين
        # ============================================================
        
        all_columns = run_query('''
            SELECT "رقم اللوحة", "اسم العمود" as "الموقع", "العدد", "الشبكة", "الحجم" 
            FROM "اعمدة انارة" 
            WHERE "المحافظة" = %s AND "الحجم" = %s
        ''', (selected_city, selected_size))
        
        # ============================================================
        # 3. تحديد الأعمدة المحجوزة في الفترات المحددة
        # ============================================================
        
        period_placeholders = ','.join([f"'{p}'" for p in selected_periods])
        booked_query = f'''
            SELECT DISTINCT "رقم اللوحة" FROM "حجوزات1" 
            WHERE "العام" = %s 
            AND "فترة الحجز" IN ({period_placeholders})
        '''
        booked_df = run_query(booked_query, (year,))
        booked_boards = booked_df['رقم اللوحة'].tolist() if booked_df is not None and not booked_df.empty else []
        
        # ============================================================
        # 4. تصفية الأعمدة المتاحة (غير المحجوزة)
        # ============================================================
        
        available_columns = all_columns[~all_columns['رقم اللوحة'].isin(booked_boards)]
        
        if available_columns.empty:
            st.warning("⚠️ لا توجد مواقع متاحة")
        else:
            # ============================================================
            # 5. عرض الشبكات المتاحة (التي تحتوي على أعمدة متاحة)
            # ============================================================
            
            # تجميع الشبكات مع عدد الأعمدة المتاحة فيها
            network_summary = available_columns.groupby('الشبكة').agg({
                'رقم اللوحة': 'count',
                'العدد': 'sum'
            }).reset_index()
            network_summary.columns = ['الشبكة', 'عدد الأعمدة', 'إجمالي اللوحات']
            
            st.markdown("**📡 الشبكات المتاحة:**")
            
            # عرض الشبكات كأزرار اختيار
            selected_network = st.selectbox(
                "اختر الشبكة:",
                network_summary['الشبكة'].tolist(),
                format_func=lambda x: f"{x} - {network_summary[network_summary['الشبكة'] == x]['عدد الأعمدة'].iloc[0]} أعمدة متاحة"
            )
            
            if selected_network is not None:
                # الأعمدة المتاحة في هذه الشبكة
                network_columns = available_columns[available_columns['الشبكة'] == selected_network]
                
                st.markdown(f"**📍 أعمدة شبكة {selected_network} المتاحة:**")
                
                # عرض الأعمدة في جدول مع خيارات اختيار فردية
                st.dataframe(
                    network_columns[['رقم اللوحة', 'الموقع', 'العدد']],
                    use_container_width=True,
                    height=200
                )
                
                # ============================================================
                # 6. خيارات الإضافة (شبكة كاملة أو أعمدة محددة)
                # ============================================================
                
                col_add1, col_add2 = st.columns(2)
                
                with col_add1:
                    # إضافة الشبكة كاملة
                    if st.button(f"📡 إضافة شبكة {selected_network} كاملة", use_container_width=True):
                        if selected_city not in st.session_state.cart:
                            st.session_state.cart[selected_city] = {}
                        
                        net_data = network_columns.copy()
                        net_data['fee_print'] = per_column_print
                        net_data['fee_display'] = per_column_display
                        
                        st.session_state.cart[selected_city][f"شبكة {selected_network} (كاملة)"] = net_data
                        st.success(f"✅ تمت إضافة شبكة {selected_network} كاملة ({len(net_data)} أعمدة)")
                        st.rerun()
                
                with col_add2:
                    # اختيار أعمدة محددة من الشبكة
                    st.write("**اختر أعمدة محددة:**")
                    
                    # استخدام multiselect لاختيار أعمدة محددة
                    selected_boards = st.multiselect(
                        "اختر الأعمدة:",
                        network_columns['رقم اللوحة'].tolist(),
                        format_func=lambda x: f"{x} - {network_columns[network_columns['رقم اللوحة'] == x]['الموقع'].iloc[0]}",
                        key=f"individual_select_{selected_network}"
                    )
                    
                    if selected_boards and st.button(f"📍 إضافة الأعمدة المحددة ({len(selected_boards)})", use_container_width=True):
                        individual_data = network_columns[network_columns['رقم اللوحة'].isin(selected_boards)].copy()
                        individual_data['fee_print'] = per_column_print
                        individual_data['fee_display'] = per_column_display
                        
                        if selected_city not in st.session_state.cart:
                            st.session_state.cart[selected_city] = {}
                        
                        st.session_state.cart[selected_city][f"أعمدة من شبكة {selected_network}"] = individual_data
                        st.success(f"✅ تمت إضافة {len(individual_data)} أعمدة محددة")
                        st.rerun()
        
        if st.session_state.cart:
            st.divider()
            st.subheader("🛒 سلة العروض")
            
            grand_total_print = 0.0
            grand_total_display = 0.0
            
            for city, items in list(st.session_state.cart.items()):
                for item_name, df_cart in list(items.items()):
                    # تحديد نوع العنصر
                    if "شبكة" in item_name:
                        icon = "📡"
                    else:
                        icon = "📍"
                    
                    with st.expander(f"{icon} {city} - {item_name}", expanded=True):
                        # عرض الأعمدة في الجدول مع إمكانية التعديل
                        edited_df = st.data_editor(
                            df_cart, 
                            key=f"edit_{city}_{item_name}", 
                            num_rows="dynamic", 
                            use_container_width=True
                        )
                        st.session_state.cart[city][item_name] = edited_df
                        
                        qty = int(edited_df['العدد'].sum())
                        fp = float(edited_df['fee_print'].iloc[0]) if 'fee_print' in edited_df.columns else per_column_print
                        fd = float(edited_df['fee_display'].iloc[0]) if 'fee_display' in edited_df.columns else per_column_display
                        
                        section_print = qty * fp
                        section_display = qty * fd
                        
                        grand_total_print += section_print
                        grand_total_display += section_display
                        
                        st.info(f"📊 العدد: {qty} | الطباعة: {section_print:.2f}$ | العرض: {section_display:.2f}$")
                        
                        if st.button("🗑️ حذف", key=f"delete_{city}_{item_name}"):
                            del st.session_state.cart[city][item_name]
                            st.rerun()
            
            st.divider()
            
            st.subheader("💰 خيارات الحسم")
            
            col_disc1, col_disc2 = st.columns([1, 2])
            with col_disc1:
                apply_discount = st.checkbox("🏷️ تطبيق حسم على أجور العرض فقط")
            with col_disc2:
                discount_percent = 0
                if apply_discount:
                    discount_percent = st.slider("نسبة الحسم (%)", min_value=1, max_value=99, value=10, step=1)
            
            if apply_discount and discount_percent > 0:
                discount_amount = grand_total_display * (discount_percent / 100)
                grand_total_display_after = grand_total_display - discount_amount
                grand_total = grand_total_print + grand_total_display_after
                
                st.info(f"""
                💰 **تفاصيل الفاتورة:**
                - إجمالي أجور الطباعة: **{grand_total_print:,.2f} $**
                - إجمالي أجور العرض (قبل الحسم): **{grand_total_display:,.2f} $**
                - حسم **{discount_percent}%**: **- {discount_amount:,.2f} $**
                - إجمالي أجور العرض (بعد الحسم): **{grand_total_display_after:,.2f} $**
                """)
            else:
                grand_total = grand_total_print + grand_total_display
                st.info(f"""
                💰 **تفاصيل الفاتورة:**
                - إجمالي أجور الطباعة: **{grand_total_print:,.2f} $**
                - إجمالي أجور العرض: **{grand_total_display:,.2f} $**
                """)
            
            st.success(f"## 💰 الإجمالي النهائي: {grand_total:,.2f} $")
            
            col_btn1, col_btn2, col_btn3, col_btn4 = st.columns(4)
            
            with col_btn1:
                if st.button("💾 حفظ كمسودة", use_container_width=True, key="save_draft"):
                    if not customer_name:
                        st.error("❌ الرجاء إدخال اسم الزبون")
                    else:
                        save_data = {"data": {c: {n: df.to_dict() for n, df in ns.items()} for c, ns in st.session_state.cart.items()}}
                        cursor = conn.cursor()
                        cursor.execute('''
                            INSERT INTO "offers_history" (client_name, cart_json, status, start_p, end_p, year, offer_date) 
                            VALUES (%s, %s, %s, %s, %s, %s, NOW())
                        ''', (customer_name, json.dumps(save_data, ensure_ascii=False), 'Pending', start_p, end_p, year))
                        conn.commit()
                        cursor.close()
                        st.success("✅ تم الحفظ كمسودة")
            
            with col_btn2:
                if is_admin():
                    if st.button("✅ تثبيت نهائي", use_container_width=True, key="confirm_booking"):
                        if not customer_name:
                            st.error("❌ الرجاء إدخال اسم الزبون")
                        else:
                            try:
                                cur = conn.cursor()
                                for city, networks in st.session_state.cart.items():
                                    for net, df in networks.items():
                                        for _, row in df.iterrows():
                                            for period in selected_periods:
                                                cur.execute('''
                                                    INSERT INTO "حجوزات1" ("رقم اللوحة", "اسم الزبون", "العام", "فترة الحجز") 
                                                    VALUES (%s, %s, %s, %s)
                                                ''', (str(row['رقم اللوحة']), customer_name, year, period))
                                
                                conn.commit()
                                st.session_state.cart = {}
                                st.success("✅ تم تثبيت الحجز بنجاح")
                                st.rerun()
                            except Exception as e:
                                conn.rollback()
                                st.error(f"❌ حدث خطأ: {str(e)}")
                else:
                    st.button("✅ تثبيت نهائي", use_container_width=True, disabled=True, key="confirm_booking_disabled")
                    st.caption("🔒 غير مسموح - فقط للمديرين")
            
            with col_btn3:
                if st.button("📝 تصدير Word", use_container_width=True, key="export_word"):
                    discount = discount_percent if apply_discount else 0
                    
                    doc = Document('template.docx') if os.path.exists('template.docx') else Document()
                    PURPLE_COLOR = "660099"
                    
                    discount_amount = grand_total_display * (discount / 100)
                    grand_total_display_after = grand_total_display - discount_amount
                    final_total = grand_total_print + grand_total_display_after
                    
                    doc.add_paragraph()
                    today_date = datetime.now().strftime("%d / %m / %Y")
                    p_date = doc.add_paragraph()
                    p_date.add_run(f"التاريخ: {today_date}")
                    _force_rtl_style(p_date)
                    doc.add_paragraph()
                    
                    p_cust = doc.add_paragraph()
                    p_cust.add_run(f"السادة شركة {customer_name} المحترمين").bold = True
                    _force_rtl_style(p_cust)
                    
                    # ✅ استخدام الصيغة المفهومة للزبون
                    start_display = format_period_for_display(start_p, take_first=False)
                    end_display = format_period_for_display(end_p, take_first=True)
                    p_stat = doc.add_paragraph()
                    p_stat.add_run(f"نقدم لكم المواقع المتاحة لعرض إعلانكم الوطني اعتباراً من {start_display} لغاية {end_display}")
                    _force_rtl_style(p_stat)
                    st.write(f"🔍 start_p: {start_p}")
                    st.write(f"🔍 end_p: {end_p}")
                    st.write(f"🔍 start_display: {start_display}")
                    st.write(f"🔍 end_display: {end_display}")
                    # ... باقي الكود (الجداول، الحسابات، إلخ)
                    
                    for city, networks in st.session_state.cart.items():
                        p_city = doc.add_paragraph()
                        p_city.add_run(f"■ محافظة {city}").bold = True
                        _force_rtl_style(p_city)
                        
                        for net, df in networks.items():
                            if df.empty:
                                continue
                            for size_info, group_df in df.groupby(['الحجم']):
                                p_size = doc.add_paragraph()
                                p_size.add_run(f"الشبكة: {net} | القياس: {size_info}").bold = True
                                _force_rtl_style(p_size)
                                
                                table = doc.add_table(rows=1, cols=2)
                                table.style = 'Table Grid'
                                set_table_rtl(table)
                                
                                hdr = table.rows[0].cells
                                hdr[0].text = "اسم الموقع (العمود)"
                                hdr[1].text = "العدد"
                                for cell in hdr:
                                    for p in cell.paragraphs:
                                        _force_rtl_style(p)
                                    tc_pr = cell._element.get_or_add_tcPr()
                                    shd = OxmlElement('w:shd')
                                    shd.set(qn('w:fill'), PURPLE_COLOR)
                                    tc_pr.append(shd)
                                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                                
                                for _, row in group_df.iterrows():
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = str(row['الموقع'])
                                    row_cells[1].text = str(row['العدد'])
                                    for cell in row_cells:
                                        for p in cell.paragraphs:
                                            _force_rtl_style(p)
                                
                                total_q = pd.to_numeric(group_df['العدد']).sum()
                                fp = float(group_df['fee_print'].iloc[0])
                                fd = float(group_df['fee_display'].iloc[0])
                                sum_print = total_q * fp
                                sum_display = total_q * fd
                                
                                p_fin = doc.add_paragraph()
                                txt = (f"إجمالي العدد: {int(total_q)} | "
                                       f"أجور الطباعة: {sum_print:,.0f}$ | "
                                       f"أجور العرض: {sum_display:,.0f}$ | "
                                       f"المجموع: {sum_print + sum_display:,.0f}$")
                                p_fin.add_run(txt).bold = True
                                _force_rtl_style(p_fin)
                    
                    doc.add_paragraph()
                    
                    if discount > 0:
                        p_discount = doc.add_paragraph()
                        p_discount.add_run(f"إجمالي أجور الطباعة: {grand_total_print:,.0f} $").bold = True
                        _force_rtl_style(p_discount)
                        
                        p_discount = doc.add_paragraph()
                        p_discount.add_run(f"إجمالي أجور العرض قبل الحسم: {grand_total_display:,.0f} $").bold = True
                        _force_rtl_style(p_discount)
                        
                        p_discount = doc.add_paragraph()
                        p_discount.add_run(f"حسم {discount}% على أجور العرض: - {discount_amount:,.0f} $").bold = True
                        _force_rtl_style(p_discount)
                        
                        p_discount = doc.add_paragraph()
                        p_discount.add_run(f"إجمالي أجور العرض بعد الحسم: {grand_total_display_after:,.0f} $").bold = True
                        _force_rtl_style(p_discount)
                    else:
                        p_total_print = doc.add_paragraph()
                        p_total_print.add_run(f"إجمالي أجور الطباعة: {grand_total_print:,.0f} $").bold = True
                        _force_rtl_style(p_total_print)
                        
                        p_total_display = doc.add_paragraph()
                        p_total_display.add_run(f"إجمالي أجور العرض: {grand_total_display:,.0f} $").bold = True
                        _force_rtl_style(p_total_display)
                    
                    doc.add_paragraph()
                    p_grand = doc.add_paragraph()
                    run_g = p_grand.add_run(f"الإجمالي النهائي للعرض: {final_total:,.0f} $")
                    run_g.bold = True
                    run_g.font.size = Pt(14)
                    run_g.font.color.rgb = RGBColor(102, 0, 153)
                    _force_rtl_style(p_grand)
                    
                    doc.add_paragraph()
                    p_note = doc.add_paragraph()
                    run_note = p_note.add_run("• ملاحظة: هذه المواقع متاحة لمدة 48 ساعة.")
                    run_note.bold = True
                    _force_rtl_style(p_note)
                    
                    target = io.BytesIO()
                    doc.save(target)
                    target.seek(0)
                    
                    st.download_button("📥 تحميل العرض", target, f"Offer_{customer_name}.docx", key="download_word")
            
            with col_btn4:
                if st.button("🔴 تفريغ السلة", use_container_width=True, key="clear_cart"):
                    st.session_state.cart = {}
                    st.rerun()
    
    except Exception as e:
        st.error(f"❌ حدث خطأ: {str(e)}")

elif page == "📋 تقرير الجرد":
    st.title("📋 التقرير التجميعي - جرد اللوحات")
    st.markdown('<div class="custom-divider"></div>', unsafe_allow_html=True)
    
    try:
        periods_df = run_query('SELECT "no", "namee" FROM "الفترة" ORDER BY "no"')
        period_names = periods_df['namee'].tolist()
        
        col1, col2, col3 = st.columns(3)
        with col1:
            from_period = st.selectbox("من فترة:", period_names, key="from_period")
        with col2:
            to_period = st.selectbox("إلى فترة:", period_names, index=len(period_names)-1, key="to_period")
        with col3:
            report_year = st.number_input("العام:", value=datetime.now().year, key="report_year")
        
        from_idx = int(periods_df[periods_df['namee'] == from_period]['no'].iloc[0])
        to_idx = int(periods_df[periods_df['namee'] == to_period]['no'].iloc[0])
        target_periods = periods_df[(periods_df['no'] >= from_idx) & (periods_df['no'] <= to_idx)]['namee'].tolist()
        
        all_boards = run_query('SELECT "رقم اللوحة", "المحافظة", "الحجم", "العدد" FROM "اعمدة انارة"')
        
        period_placeholders = ','.join([f"'{p}'" for p in target_periods])
        booked_query = f'''
            SELECT DISTINCT "رقم اللوحة" 
            FROM "حجوزات1" 
            WHERE "العام" = %s 
            AND "فترة الحجز" IN ({period_placeholders})
        '''
        booked_in_period = run_query(booked_query, (report_year,))['رقم اللوحة'].tolist()
        
        all_boards['الحالة'] = all_boards['رقم اللوحة'].apply(lambda x: 'محجوز' if x in booked_in_period else 'متاح')
        
        total_sites = len(all_boards)
        booked_sites = len(booked_in_period)
        available_sites = total_sites - booked_sites
        total_boards_count = all_boards['العدد'].sum()
        booked_boards_count = all_boards[all_boards['الحالة'] == 'محجوز']['العدد'].sum()
        available_boards_count = total_boards_count - booked_boards_count
        
        cols = st.columns(4)
        metrics_data = [
            ("🏢 إجمالي المواقع", total_sites, "🗺️", "primary"),
            ("🔴 المواقع المحجوزة", booked_sites, "📌", "danger"),
            ("🟢 المواقع المتاحة", available_sites, "✅", "success"),
            ("📈 نسبة الإشغال", f"{(booked_sites/total_sites*100):.1f}%", "📊", "warning")
        ]
        
        for idx, (title, value, icon, color) in enumerate(metrics_data):
            with cols[idx]:
                st.markdown(create_metric_card_3d(title, value, icon, color), unsafe_allow_html=True)
        
        st.divider()
        
        col_chart1, col_chart2 = st.columns(2)
        
        with col_chart1:
            fig_pie = go.Figure(data=[go.Pie(
                labels=['محجوز', 'متاح'],
                values=[booked_boards_count, available_boards_count],
                hole=0.4,
                marker_colors=['#dc2626', '#22c55e'],
                textinfo='percent+label'
            )])
            fig_pie.update_layout(title="نسبة إشغال الأعمدة", height=400)
            st.plotly_chart(fig_pie, use_container_width=True)
        
        with col_chart2:
            city_data = []
            for city in all_boards['المحافظة'].unique():
                city_df = all_boards[all_boards['المحافظة'] == city]
                city_total = city_df['العدد'].sum()
                city_booked = city_df[city_df['الحالة'] == 'محجوز']['العدد'].sum()
                city_data.append({
                    'المحافظة': city,
                    'نسبة الإشغال': (city_booked / city_total * 100) if city_total > 0 else 0
                })
            
            city_df = pd.DataFrame(city_data)
            fig_bar = px.bar(city_df, x='المحافظة', y='نسبة الإشغال', 
                           color='نسبة الإشغال', color_continuous_scale='RdYlGn')
            fig_bar.update_layout(height=400)
            st.plotly_chart(fig_bar, use_container_width=True)
        
        st.divider()
        
        st.subheader("📋 تفصيل حسب المحافظة")
        city_details = []
        for city in all_boards['المحافظة'].unique():
            city_df = all_boards[all_boards['المحافظة'] == city]
            city_total = city_df['العدد'].sum()
            city_booked = city_df[city_df['الحالة'] == 'محجوز']['العدد'].sum()
            city_details.append({
                'المحافظة': city,
                'الإجمالي': int(city_total),
                'محجوز': int(city_booked),
                'متاح': int(city_total - city_booked),
                'نسبة الإشغال': f"{(city_booked/city_total*100):.1f}%" if city_total > 0 else "0%"
            })
        
        st.dataframe(pd.DataFrame(city_details), use_container_width=True)
        
        st.divider()
        csv_data = all_boards.to_csv(index=False, encoding='utf-8-sig')
        st.download_button("📊 تصدير إلى CSV", csv_data, f"Inventory_Report_{report_year}.csv", "text/csv", use_container_width=True)
        
    except Exception as e:
        st.error(f"حدث خطأ في التقرير: {str(e)}")

elif page == "📅 تقرير التوفر الشهري":
    st.title("📋 تقرير الأعمدة المتاحة")
    st.info("📌 يعرض هذا التقرير الأعمدة المتاحة حالياً أو التي ستصبح متاحة بعد تاريخ محدد")
    
    current_year = date.today().year
    today = date.today()
    
    col_filter1, col_filter2 = st.columns(2)
    with col_filter1:
        show_all = st.checkbox("📅 عرض جميع الأعمدة المتاحة حالياً", value=True)
    with col_filter2:
        future_date = st.date_input("🗓️ عرض الأعمدة التي ستصبح متاحة بعد تاريخ", value=today + timedelta(days=7))
    
    notes = st.text_area("📝 ملاحظات (تظهر في نهاية التقرير)", placeholder="أضف ملاحظاتك هنا...", height=100)
    
    if st.button("🚀 تشغيل التقرير", use_container_width=True, type="primary"):
        with st.spinner("جاري إنشاء التقرير..."):
            all_columns = run_query('SELECT "رقم اللوحة", "اسم العمود", "المحافظة", "الشبكة", "الحجم", "العدد" FROM "اعمدة انارة"')
            
            if show_all:
                bookings_query = 'SELECT DISTINCT "رقم اللوحة" FROM "حجوزات1" WHERE "العام" = %s'
                booked_df = run_query(bookings_query, (current_year,))
            else:
                bookings_query = '''
                    SELECT DISTINCT "رقم اللوحة" FROM "حجوزات1" 
                    WHERE "العام" = %s
                    AND ("تاريخ النهاية" >= %s OR "فترة الحجز" IS NOT NULL)
                '''
                booked_df = run_query(bookings_query, (current_year, future_date))
            
            booked_boards = booked_df['رقم اللوحة'].tolist() if booked_df is not None and not booked_df.empty else []
            
            available_df = all_columns[~all_columns['رقم اللوحة'].isin(booked_boards)]
            total_available = len(available_df)
            total_boards_count = available_df['العدد'].sum() if 'العدد' in available_df.columns else total_available
            
            st.success(f"✅ {total_available} موقعاً ({int(total_boards_count)} لوحة) متاحة")
            
            st.subheader("📊 ملخص حسب المحافظة")
            summary = available_df.groupby('المحافظة').agg({
                'رقم اللوحة': 'count',
                'العدد': 'sum'
            }).rename(columns={'رقم اللوحة': 'عدد المواقع', 'العدد': 'عدد اللوحات'})
            st.dataframe(summary, use_container_width=True)
            
            st.subheader("📋 قائمة الأعمدة المتاحة")
            st.dataframe(available_df[['رقم اللوحة', 'اسم العمود', 'المحافظة', 'الشبكة', 'الحجم', 'العدد']], use_container_width=True, height=400)
            
            csv_data = available_df.to_csv(index=False, encoding='utf-8-sig')
            st.download_button("📥 تحميل التقرير (CSV)", csv_data, f"available_columns_{date.today().strftime('%Y%m%d')}.csv", "text/csv", use_container_width=True)

elif page == "🗺️ تقرير جميع المواقع":
    st.title("🗺️ تقرير جميع المواقع والأعمدة")
    st.info("📌 يعرض هذا التقرير جميع المواقع والأعمدة في النظام")
    
    # جلب البيانات
    all_columns = run_query('SELECT "رقم اللوحة", "اسم العمود", "المحافظة", "الشبكة", "الحجم", "العدد" FROM "اعمدة انارة" ORDER BY "المحافظة", "الشبكة"')
    
    # تشخيص سريع
    st.write(f"**Debug:** عدد السجلات = {len(all_columns) if all_columns is not None else 0}")
    if all_columns is not None and not all_columns.empty:
        st.write(f"**Debug:** الأعمدة الموجودة: {all_columns.columns.tolist()}")
    
    if all_columns is None or all_columns.empty:
        st.warning("⚠️ لا توجد بيانات في جدول أعمدة الإنارة")
        st.stop()
    
    # إحصائيات سريعة
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("إجمالي المواقع", len(all_columns))
    with col2:
        st.metric("إجمالي الأعمدة", int(all_columns['العدد'].sum()) if 'العدد' in all_columns.columns else len(all_columns))
    with col3:
        st.metric("عدد المحافظات", all_columns['المحافظة'].nunique() if 'المحافظة' in all_columns.columns else 0)
    
    st.divider()
    
    # عرض الجدول كاملاً أولاً (للتأكد من وجود بيانات)
    st.subheader("📋 جميع البيانات (جدول كامل)")
    st.dataframe(all_columns, use_container_width=True)
    
    st.divider()
    
    # عرض البيانات بشكل منظم حسب المحافظة
    st.subheader("📋 تفصيل حسب المحافظة")
    
    for city in sorted(all_columns['المحافظة'].unique()):
        city_df = all_columns[all_columns['المحافظة'] == city]
        
        with st.expander(f"📍 محافظة {city} ({len(city_df)} موقع - {city_df['العدد'].sum()} لوحة)"):
            
            # عرض جميع مواقع المحافظة
            st.dataframe(city_df[['رقم اللوحة', 'اسم العمود', 'الشبكة', 'الحجم', 'العدد']], use_container_width=True)
            
            # تفصيل حسب الشبكة
            if 'الشبكة' in city_df.columns:
                st.write("**📡 تفصيل حسب الشبكة:**")
                network_summary = city_df.groupby('الشبكة').agg({
                    'رقم اللوحة': 'count',
                    'العدد': 'sum'
                }).rename(columns={'رقم اللوحة': 'عدد المواقع', 'العدد': 'عدد الأعمدة'})
                st.dataframe(network_summary, use_container_width=True)
    
    # تصدير
    st.divider()
    csv_data = all_columns.to_csv(index=False, encoding='utf-8-sig')
    st.download_button("📊 تصدير CSV", csv_data, f"full_report_{date.today().strftime('%Y%m%d')}.csv", "text/csv", use_container_width=True)

elif page == "📐 تقرير تجميعي حسب الحجوم":
    st.title("📐 تقرير تجميعي حسب الحجوم")
    st.info("📌 يعرض هذا التقرير توزع اللوحات حسب الحجوم المقسمة إلى ثلاث مجموعات")
    
    all_columns = run_query('SELECT "رقم اللوحة", "اسم العمود", "المحافظة", "الشبكة", "الحجم", "العدد" FROM "اعمدة انارة" ORDER BY "المحافظة", "الشبكة"')
    
    if all_columns is None or all_columns.empty:
        st.warning("⚠️ لا توجد بيانات في جدول الأعمدة")
        st.stop()
    
    group1_sizes = ['3*6', '3x6', '3 × 6']
    group2_sizes = ['2*1', '2x1', '2 × 1', '125*185', '125x185', '125 × 185']
    
    def classify_size(size):
        size_str = str(size).strip()
        if size_str in group1_sizes or size_str.replace(' ', '') in ['3*6', '3x6']:
            return 'المجموعة الأولى: حجم 3×6'
        elif size_str in group2_sizes or size_str.replace(' ', '') in ['2*1', '2x1', '125*185', '125x185']:
            return 'المجموعة الثانية: حجمي 2×1 و 125×185'
        else:
            return 'المجموعة الثالثة: باقي الحجوم'
    
    all_columns['المجموعة'] = all_columns['الحجم'].apply(classify_size)
    
    cols = st.columns(3)
    with cols[0]:
        st.markdown(create_metric_card_3d("إجمالي الأعمدة", int(all_columns['العدد'].sum()), "📌", "primary"), unsafe_allow_html=True)
    with cols[1]:
        st.markdown(create_metric_card_3d("إجمالي المواقع", len(all_columns), "🗺️", "success"), unsafe_allow_html=True)
    with cols[2]:
        st.markdown(create_metric_card_3d("عدد الأحجام", all_columns['الحجم'].nunique(), "📏", "warning"), unsafe_allow_html=True)
    
    st.divider()
    
    st.subheader("📊 ملخص المجموعات")
    group_summary = all_columns.groupby('المجموعة').agg({
        'رقم اللوحة': 'count',
        'العدد': 'sum'
    }).rename(columns={'رقم اللوحة': 'عدد المواقع', 'العدد': 'عدد الأعمدة'})
    group_summary['عدد الأعمدة'] = group_summary['عدد الأعمدة'].astype(int)
    st.dataframe(group_summary, use_container_width=True)
    
    st.divider()
    
    for group_name in ['المجموعة الأولى: حجم 3×6', 'المجموعة الثانية: حجمي 2×1 و 125×185', 'المجموعة الثالثة: باقي الحجوم']:
        group_df = all_columns[all_columns['المجموعة'] == group_name]
        if not group_df.empty:
            with st.expander(f"📌 {group_name} - {len(group_df)} موقع - {int(group_df['العدد'].sum())} عمود", expanded=False):
                st.subheader("📍 توزع حسب المحافظة")
                city_summary = group_df.groupby('المحافظة').agg({
                    'رقم اللوحة': 'count',
                    'العدد': 'sum'
                }).rename(columns={'رقم اللوحة': 'عدد المواقع', 'العدد': 'عدد الأعمدة'})
                st.dataframe(city_summary, use_container_width=True)
                
                st.subheader("📋 قائمة المواقع")
                st.dataframe(group_df[['رقم اللوحة', 'اسم العمود', 'المحافظة', 'الشبكة', 'الحجم', 'العدد']], use_container_width=True)
    
    st.divider()
    
    csv_data = all_columns.to_csv(index=False, encoding='utf-8-sig')
    st.download_button("📊 تصدير التقرير كاملاً (CSV)", csv_data, f"grouped_report_{date.today().strftime('%Y%m%d')}.csv", "text/csv", use_container_width=True)



#=============
# التحكم
#=============
elif page == "⚙️ الإعدادات":
    if not is_admin():
        st.error("⛔ هذه الصفحة مخصصة للمديرين فقط")
        st.stop()
    
    st.title("⚙️ إعدادات النظام - إدارة البيانات")
    st.warning("⚠️ تحذير: تعديل هذه البيانات يؤثر مباشرة على النظام. يرجى الحذر.")
    
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) FROM \"اعمدة انارة\"")
    boards_count = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM \"حجوزات1\"")
    bookings_count = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM \"اسماء الرسم\"")
    fees_count = cursor.fetchone()[0]
    cursor.close()
    
    cols = st.columns(3)
    with cols[0]:
        st.markdown(create_metric_card_3d("أعمدة الإنارة", boards_count, "🗺️", "primary"), unsafe_allow_html=True)
    with cols[1]:
        st.markdown(create_metric_card_3d("الحجوزات", bookings_count, "📅", "success"), unsafe_allow_html=True)
    with cols[2]:
        st.markdown(create_metric_card_3d("أجور الرسم", fees_count, "💰", "warning"), unsafe_allow_html=True)
    
    st.divider()
    
    tab1, tab2, tab3, tab4 = st.tabs(["🗄️ أعمدة الإنارة", "📅 سجل الحجوزات", "💰 أجور الرسم", "👥 المستخدمين"])
    
    with tab1:
        st.subheader("إدارة بيانات أعمدة الإنارة")
        df_boards = run_query('SELECT * FROM "اعمدة انارة" ORDER BY "المحافظة", "الشبكة"')
        edited_boards = st.data_editor(df_boards, num_rows="dynamic", key="edit_boards", use_container_width=True)
        if st.button("💾 حفظ أعمدة الإنارة", key="save_boards", use_container_width=True):
            cursor = conn.cursor()
            cursor.execute("DELETE FROM \"اعمدة انارة\"")
            for _, row in edited_boards.iterrows():
                cursor.execute('''
                    INSERT INTO "اعمدة انارة" ("رقم اللوحة", "اسم العمود", "المحافظة", "الشبكة", "الحجم", "العدد", "Latitude", "Longitude")
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                ''', (row['رقم اللوحة'], row['اسم العمود'], row['المحافظة'], row['الشبكة'], row['الحجم'], row['العدد'], 
                      row.get('Latitude'), row.get('Longitude')))
            conn.commit()
            cursor.close()
            st.success("✅ تم تحديث أعمدة الإنارة")
            st.rerun()
    
    with tab2:
        st.subheader("إدارة سجل الحجوزات")
        df_bookings = run_query('SELECT * FROM "حجوزات1"')
        edited_bookings = st.data_editor(df_bookings, num_rows="dynamic", key="edit_bookings", use_container_width=True)
        if st.button("💾 حفظ سجل الحجوزات", key="save_bookings", use_container_width=True):
            cursor = conn.cursor()
            cursor.execute("DELETE FROM \"حجوزات1\"")
            for _, row in edited_bookings.iterrows():
                cursor.execute('''
                    INSERT INTO "حجوزات1" ("رقم اللوحة", "اسم الزبون", "العام", "فترة الحجز", "تاريخ النهاية")
                    VALUES (%s, %s, %s, %s, %s)
                ''', (row['رقم اللوحة'], row['اسم الزبون'], row['العام'], row['فترة الحجز'], row.get('تاريخ النهاية')))
            conn.commit()
            cursor.close()
            st.success("✅ تم تحديث سجل الحجوزات")
            st.rerun()
    
    with tab3:
        st.subheader("إدارة أجور الرسم")
        st.info("💡 أضف 'اجور الطباعة عادي' و 'اجور الطباعة سكوتش' و 'اجور العرض شهري' و 'اجور العرض اجنبي شهري'")
        df_fees = run_query('SELECT * FROM "اسماء الرسم"')
        edited_fees = st.data_editor(df_fees, num_rows="dynamic", key="edit_fees", use_container_width=True)
        if st.button("💾 حفظ أجور الرسم", key="save_fees", use_container_width=True):
            cursor = conn.cursor()
            cursor.execute("DELETE FROM \"اسماء الرسم\"")
            for _, row in edited_fees.iterrows():
                cursor.execute('''
                    INSERT INTO "اسماء الرسم" ("اسم الرسم", "الحجم", "اجرة الرسم")
                    VALUES (%s, %s, %s)
                ''', (row['اسم الرسم'], row['الحجم'], row['اجرة الرسم']))
            conn.commit()
            cursor.close()
            st.success("✅ تم تحديث أجور الرسم")
            st.rerun()
    
    with tab4:
        st.subheader("👥 إدارة المستخدمين")
        df_users = run_query('SELECT id, username, role, full_name, created_at FROM users')
        edited_users = st.data_editor(df_users, num_rows="dynamic", key="edit_users", use_container_width=True)
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 حفظ المستخدمين", key="save_users", use_container_width=True):
                cursor = conn.cursor()
                for _, row in edited_users.iterrows():
                    cursor.execute('''
                        UPDATE users SET username=%s, role=%s, full_name=%s WHERE id=%s
                    ''', (row['username'], row['role'], row['full_name'], row['id']))
                conn.commit()
                cursor.close()
                st.success("✅ تم تحديث المستخدمين")
                st.rerun()
        
        with col2:
            with st.expander("➕ إضافة مستخدم جديد"):
                new_username = st.text_input("اسم المستخدم")
                new_password = st.text_input("كلمة المرور", type="password")
                new_role = st.selectbox("الدور", ["admin", "employee"])
                new_full_name = st.text_input("الاسم الكامل")
                if st.button("إضافة مستخدم", use_container_width=True):
                    cursor = conn.cursor()
                    try:
                        cursor.execute('''
                            INSERT INTO users (username, password, role, full_name, created_at)
                            VALUES (%s, %s, %s, %s, NOW())
                        ''', (new_username, new_password, new_role, new_full_name))
                        conn.commit()
                        cursor.close()
                        st.success("✅ تم إضافة المستخدم")
                        st.rerun()
                    except Exception as e:
                        st.error(f"خطأ: {e}")

# =============
# صفحة الإدخال اليومي (النسخة المعدلة)
# =============

elif page == "📝 الإدخال اليومي":
    # التحقق من الصلاحيات
    if not is_authenticated():
        st.error("⛔ يرجى تسجيل الدخول أولاً")
        st.stop()
    
    # الهيدر
    st.markdown("""
    <div class="main-header">
        <h1>📝 نظام الإدخال اليومي</h1>
        <p>إدارة الحجوزات والأعمدة بكفاءة وسهولة</p>
    </div>
    """, unsafe_allow_html=True)
    
    # معلومات المستخدم الحالي
    user_info = get_current_user()
    if user_info:
        st.sidebar.markdown(f"""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                    padding: 1rem; border-radius: 10px; color: white; margin-bottom: 1rem;">
            <strong>👤 {user_info.get('full_name', 'مستخدم')}</strong><br>
            <small style="opacity: 0.8;">@ {user_info.get('username', '')} • {user_info.get('role', 'employee')}</small>
        </div>
        """, unsafe_allow_html=True)
    
    # جلب قائمة الحجوم من قاعدة البيانات
    df_sizes = run_query('SELECT DISTINCT "الحجم" FROM "اعمدة انارة" WHERE "الحجم" IS NOT NULL ORDER BY "الحجم"')
    size_options = df_sizes['الحجم'].tolist() if not df_sizes.empty else []
    
    # خيارات الجاهزية (ثابتة)
    readiness_options = ['جاهز', 'بحاجة صيانة', 'خارج الخدمة']
    
    # تبويبات الإدخال
    if user_info and user_info.get('role') == 'admin':
        tabs = st.tabs(["📅 حجز جديد", "🗺️ إضافة عمود", "📊 عرض الحجوزات", "⚡ إجراءات سريعة"])
    else:
        tabs = st.tabs(["📅 حجز جديد", "📊 عرض حجوزاتي"])
        
    #===============تبويب حجز جديد======================
    with tabs[0]:
        st.markdown('<div class="input-card">', unsafe_allow_html=True)
        st.subheader("🆕 إنشاء حجز جديد")
        
        # تهيئة السلة في session_state
        if 'booking_cart' not in st.session_state:
            st.session_state.booking_cart = []
        
        # ============================================================
        # الفلاتر خارج الـ form
        # ============================================================
        
        col_filters = st.columns(3)
        with col_filters[0]:
            df_cities = run_query('SELECT DISTINCT "المحافظة" FROM "اعمدة انارة" ORDER BY "المحافظة"')
            city_options = df_cities['المحافظة'].tolist() if not df_cities.empty else []
            selected_city = st.selectbox("📍 المحافظة", city_options, key="city_filter")
        
        with col_filters[1]:
            if selected_city:
                try:
                    conn_local = get_connection()
                    df_networks = pd.read_sql_query("""
                        SELECT DISTINCT "الشبكة" 
                        FROM "اعمدة انارة" 
                        WHERE "المحافظة" = %s 
                        AND "الشبكة" IS NOT NULL
                        ORDER BY "الشبكة"
                    """, conn_local, params=(selected_city,))
                    conn_local.close()
                    network_options = df_networks['الشبكة'].tolist() if not df_networks.empty else []
                    
                except Exception as e:
                    st.error(f"❌ {e}")
                    network_options = []
            else:
                network_options = []
            
            selected_network = st.selectbox("🌐 الشبكة", ["جميع الشبكات"] + network_options, key="network_filter")
        
        # ============================================================
        # جلب اللوحات حسب الفلاتر
        # ============================================================
        
        if selected_city:
            try:
                conn_local = get_connection()
                if selected_network and selected_network != "جميع الشبكات":
                    boards_query = """
                        SELECT "رقم اللوحة", "اسم العمود", "الشبكة", "الحجم", "الجاهزية"
                        FROM "اعمدة انارة" 
                        WHERE "المحافظة" = %s 
                        AND "الشبكة" = %s
                        ORDER BY "رقم اللوحة"
                    """
                    df_boards = pd.read_sql_query(boards_query, conn_local, params=(selected_city, selected_network))
                else:
                    boards_query = """
                        SELECT "رقم اللوحة", "اسم العمود", "الشبكة", "الحجم", "الجاهزية"
                        FROM "اعمدة انارة" 
                        WHERE "المحافظة" = %s
                        ORDER BY "رقم اللوحة"
                    """
                    df_boards = pd.read_sql_query(boards_query, conn_local, params=(selected_city,))
                conn_local.close()
            except Exception as e:
                st.error(f"❌ خطأ في جلب اللوحات: {e}")
                df_boards = pd.DataFrame()
        else:
            df_boards = pd.DataFrame()
        
        # ============================================================
        # اختيار اللوحات
        # ============================================================
        
        if not df_boards.empty:
            st.caption(f"📊 عدد الأعمدة: {len(df_boards)}")
            board_options = df_boards.apply(
                lambda row: f"{row['رقم اللوحة']} - {row['اسم العمود']} ({row['الحجم']}) - {row.get('الجاهزية', 'جاهز')}", 
                axis=1
            ).tolist()
            
            # ✅ استخدام st.columns لعرض زر الإضافة بجانب القائمة
            col_select, col_btn = st.columns([3, 1])
            
            with col_select:
                selected_boards = st.multiselect(
                    "🏷️ اختيار اللوحات",
                    board_options,
                    key="boards_select_outside",
                    placeholder="اختر لوحة أو أكثر..."
                )
            
            with col_btn:
                st.write("")
                st.write("")
                if st.button("➕ إضافة إلى السلة", key="add_to_cart_btn", use_container_width=True):
                    if selected_boards:
                        new_boards = [b.split(' - ')[0] for b in selected_boards]
                        st.session_state.booking_cart.extend(new_boards)
                        st.success(f"✅ تم إضافة {len(new_boards)} لوحة")
                        st.rerun()
                    else:
                        st.warning("⚠️ يرجى اختيار لوحة أولاً")
        else:
            st.warning("⚠️ لا توجد لوحات في هذه المحافظة")
        
        # ============================================================
        # عرض سلة اللوحات المختارة
        # ============================================================
        
        if st.session_state.booking_cart:
            st.divider()
            st.subheader(f"🛒 سلة اللوحات المختارة ({len(st.session_state.booking_cart)})")
            
            # عرض اللوحات في السلة
            cart_df = pd.DataFrame(st.session_state.booking_cart, columns=['رقم اللوحة'])
            st.dataframe(cart_df, use_container_width=True, height=150)
            
            # زر تفريغ السلة
            if st.button("🗑️ تفريغ السلة", key="clear_cart_btn"):
                st.session_state.booking_cart = []
                st.rerun()
        
        # ============================================================
        # نموذج الحجز (فقط للحفظ)
        # ============================================================
        
        with st.form("booking_save_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            
            with col1:
                customer_name = st.text_input("👤 اسم الزبون", placeholder="أدخل اسم الزبون كاملاً")
                year = st.number_input("📅 العام", min_value=2020, max_value=2030, value=2026, step=1)
            
            with col2:
                today = datetime.now().date()
                booking_start = st.date_input("📆 بداية الحجز", value=today, min_value=today)
                booking_end = st.date_input("📆 نهاية الحجز", value=today + timedelta(days=30), min_value=booking_start)
                board_type = st.selectbox("📋 نوع اللوحة", ["عادية", "سكوتش"])
            
            notes = st.text_area("📝 ملاحظات", placeholder="أي معلومات إضافية...", height=80)
            
            col5, col6 = st.columns(2)
            with col5:
                phone = st.text_input("📞 الهاتف", placeholder="05xxxxxxxx")
            with col6:
                email = st.text_input("✉️ البريد", placeholder="example@email.com")
            
            if st.session_state.booking_cart:
                st.info(f"📋 عدد اللوحات المراد حجزها: {len(st.session_state.booking_cart)}")
            
            submitted = st.form_submit_button("💾 حفظ الحجز", use_container_width=True)
            
            if submitted:
                if not st.session_state.booking_cart:
                    st.error("⚠️ يرجى إضافة لوحات إلى السلة أولاً")
                elif not customer_name:
                    st.error("⚠️ يرجى إدخال اسم الزبون")
                else:
                    start_str = booking_start.strftime("%Y-%m-%d")
                    end_str = booking_end.strftime("%Y-%m-%d")
                    
                    try:
                        cursor = conn.cursor()
                        inserted = 0
                        
                        for board_number in st.session_state.booking_cart:
                            cursor.execute('''
                                INSERT INTO "حجوزات1" 
                                ("رقم اللوحة", "اسم الزبون", "العام", "فترة الحجز", "تاريخ النهاية", 
                                 "نوع اللوحة", "ملاحظات", "الهاتف", "البريد", "تاريخ الانشاء")
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                            ''', (board_number, customer_name, year, start_str, end_str,
                                  board_type, notes, phone, email))
                            inserted += 1
                        
                        conn.commit()
                        cursor.close()
                        
                        # تفريغ السلة بعد الحجز
                        st.session_state.booking_cart = []
                        
                        st.success(f"✅ تم إنشاء {inserted} حجز بنجاح!")
                        st.balloons()
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"❌ حدث خطأ: {str(e)}")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    
    # ========== تبويب إضافة عمود (للمدير فقط) ==========
    if len(tabs) > 1 and user_info and user_info.get('role') == 'admin':
        with tabs[1]:
            st.markdown('<div class="input-card">', unsafe_allow_html=True)
            st.subheader("🗺️ إضافة عمود إنارة جديد")
            
            with st.form("new_board_form", clear_on_submit=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    board_number = st.text_input("🔢 رقم اللوحة", placeholder="مثال: B001")
                    board_name = st.text_input("🏷️ اسم العمود", placeholder="مثال: شارع الملك فهد")
                    governorate = st.selectbox(
                        "📍 المحافظة",
                        ["الرياض", "جدة", "مكة", "المدينة", "الدمام", "الخبر", "الظهران", "أخرى"]
                    )
                
                with col2:
                    network = st.text_input("🌐 الشبكة", placeholder="اسم الشبكة")
                    
                    # ✅ استخدام قائمة الحجوم من قاعدة البيانات
                    board_size = st.selectbox(
                        "📐 الحجم",
                        size_options if size_options else ["صغير", "متوسط", "كبير"],
                        help="اختر حجم اللوحة من القائمة"
                    )
                    
                    # ✅ خيارات الجاهزية
                    readiness = st.selectbox(
                        "🔧 الجاهزية",
                        readiness_options,
                        help="حالة اللوحة الحالية"
                    )
                    
                    quantity = st.number_input("🔢 العدد", min_value=1, value=1)
                
                # الموقع الجغرافي
                st.markdown("---")
                st.markdown("📍 **الموقع الجغرافي (اختياري)**")
                col3, col4 = st.columns(2)
                with col3:
                    latitude = st.number_input("Latitude", format="%.6f", value=0.0)
                with col4:
                    longitude = st.number_input("Longitude", format="%.6f", value=0.0)
                
                submitted = st.form_submit_button("➕ إضافة العمود", use_container_width=True)
                
                if submitted:
                    if not board_number or not board_name:
                        st.error("⚠️ يرجى إدخال رقم اللوحة واسم العمود")
                    else:
                        try:
                            cursor = conn.cursor()
                            cursor.execute('''
                                INSERT INTO "اعمدة انارة" 
                                ("رقم اللوحة", "اسم العمود", "المحافظة", "الشبكة", "الحجم", "العدد", "Latitude", "Longitude", "الجاهزية")
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                            ''', (board_number, board_name, governorate, network, board_size, quantity, latitude, longitude, readiness))
                            conn.commit()
                            cursor.close()
                            
                            st.success("✅ تم إضافة العمود بنجاح!")
                            st.balloons()
                            
                        except Exception as e:
                            st.error(f"❌ حدث خطأ: {str(e)}")
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    # ========== تبويب عرض الحجوزات ==========
    # تحديد index التبويب الصحيح
    if user_info and user_info.get('role') == 'admin':
        tab_index = 2  # التبويب الثالث (0,1,2)
    else:
        tab_index = 1  # التبويب الثاني (0,1)
    
    with tabs[tab_index]:
        st.markdown('<div class="input-card">', unsafe_allow_html=True)
        
        if user_info and user_info.get('role') == 'admin':
            st.subheader("📊 جميع الحجوزات")
            df = run_query('SELECT * FROM "حجوزات1" ORDER BY "TimeOfTask" DESC NULLS LAST')
        else:
            st.subheader("📊 حجوزاتي")
            df = run_query(f'''
                SELECT * FROM "حجوزات1" 
                WHERE "اسم الزبون" LIKE '%{user_info.get("full_name", "")}%'
                ORDER BY "TimeOfTask" DESC NULLS LAST
            ''')
        
        if not df.empty:
            # فلترة وتصفية
            col_filter1, col_filter2, col_filter3 = st.columns(3)
            with col_filter1:
                search = st.text_input("🔍 بحث", placeholder="بحث بالزبون أو اللوحة...")
            with col_filter2:
                status_filter = st.selectbox("📌 الحالة", ["الكل", "نشط", "منتهي", "قادم"])
            with col_filter3:
                date_filter = st.date_input("📅 من تاريخ")
            
            if search:
                df = df[df['اسم الزبون'].str.contains(search, case=False) | 
                       df['رقم اللوحة'].str.contains(search, case=False)]
            
            display_cols = ['رقم اللوحة', 'اسم الزبون', 'فترة الحجز', 'تاريخ النهاية', 'نوع اللوحة']
            available_cols = [col for col in display_cols if col in df.columns]
            
            st.dataframe(
                df[available_cols],
                use_container_width=True,
                height=400,
                column_config={
                    "رقم اللوحة": "🏷️ رقم اللوحة",
                    "اسم الزبون": "👤 الزبون",
                    "فترة الحجز": "📅 بداية الحجز",
                    "تاريخ النهاية": "📅 نهاية الحجز",
                    "نوع اللوحة": "📋 النوع"
                }
            )
            
            # إحصائيات سريعة
            col_stat1, col_stat2, col_stat3 = st.columns(3)
            with col_stat1:
                st.metric("📊 إجمالي الحجوزات", len(df))
            with col_stat2:
                active = len(df[pd.to_datetime(df['تاريخ النهاية']) >= datetime.now()])
                st.metric("🟢 حجوزات نشطة", active)
            with col_stat3:
                expiring_soon = len(df[pd.to_datetime(df['تاريخ النهاية']) <= datetime.now() + timedelta(days=7)])
                st.metric("⏰ تنتهي قريباً", expiring_soon)
        else:
            st.info("📭 لا توجد حجوزات لعرضها")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # ========== تبويب الإجراءات السريعة (للمدير فقط) ==========
    if len(tabs) > 3 and user_info and user_info.get('role') == 'admin':
        with tabs[3]:
            st.markdown('<div class="input-card">', unsafe_allow_html=True)
            st.subheader("⚡ إجراءات سريعة")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### 📥 استيراد بيانات")
                uploaded_file = st.file_uploader("رفع ملف Excel", type=['xlsx', 'xls'])
                if uploaded_file:
                    if st.button("📤 استيراد"):
                        try:
                            df_import = pd.read_excel(uploaded_file)
                            st.success(f"✅ تم استيراد {len(df_import)} سجل")
                            st.dataframe(df_import.head())
                        except Exception as e:
                            st.error(f"❌ خطأ: {str(e)}")
            
            with col2:
                st.markdown("#### 📤 تصدير بيانات")
                if st.button("📥 تصدير الحجوزات"):
                    df_export = run_query('SELECT * FROM "حجوزات1"')
                    if not df_export.empty:
                        csv = df_export.to_csv(index=False)
                        st.download_button(
                            label="⬇️ تحميل CSV",
                            data=csv,
                            file_name=f"الحجوزات_{datetime.now().strftime('%Y%m%d')}.csv",
                            mime="text/csv"
                        )
            
            st.markdown("---")
            st.markdown("#### 🗑️ إدارة البيانات")
            col3, col4 = st.columns(2)
            
            with col3:
                if st.button("🧹 تنظيف الحجوزات المنتهية", use_container_width=True):
                    try:
                        cursor = conn.cursor()
                        cursor.execute('''
                            DELETE FROM "حجوزات1" 
                            WHERE "تاريخ النهاية" < CURRENT_DATE
                        ''')
                        deleted_count = cursor.rowcount
                        conn.commit()
                        cursor.close()
                        st.success(f"✅ تم حذف {deleted_count} حجز منتهي")
                    except Exception as e:
                        st.error(f"❌ خطأ: {str(e)}")
            
            with col4:
                if st.button("📊 تقرير الأداء", use_container_width=True):
                    df_bookings = run_query('SELECT * FROM "حجوزات1"')
                    if not df_bookings.empty:
                        st.info(f"""
                        📈 **إحصائيات الأداء:**
                        - إجمالي الحجوزات: {len(df_bookings)}
                        - أكثر العمود طلباً: {df_bookings['رقم اللوحة'].mode().iloc[0] if 'رقم اللوحة' in df_bookings.columns and not df_bookings['رقم اللوحة'].empty else 'N/A'}
                        """)
            
            st.markdown('</div>', unsafe_allow_html=True)



# ============================
# كتالوج عام
# ============================

elif page == "📋 كتالوج عام":
    st.title("📋 كتالوج اللوحات المتاحة")
    st.markdown('<div class="custom-divider"></div>', unsafe_allow_html=True)
    
    st.info("ℹ️ هذا الكتالوج يعرض جميع اللوحات المتاحة حالياً - بدون أسعار - للعرض العام")

    
    # تهيئة session_state
    if 'catalog_selected_boards' not in st.session_state:
        st.session_state.catalog_selected_boards = {}  # {city: [list of boards]}
    if 'catalog_cities_added' not in st.session_state:
        st.session_state.catalog_cities_added = []  # list of cities added
    
    try:
        # ============================================================
        # 1. اختيار المحافظة والحجم والفترة
        # ============================================================
        
        draw_df = run_query('SELECT * FROM "اسماء الرسم"')
        
        col1, col2 = st.columns(2)
        with col1:
            cities = run_query('SELECT DISTINCT "المحافظة" FROM "اعمدة انارة"')['المحافظة'].tolist()
            selected_city = st.selectbox("📍 اختر المحافظة:", cities)
        
        with col2:
            if not draw_df.empty:
                selected_size = st.selectbox("📏 قياس اللوحة:", draw_df['الحجم'].unique().tolist())
            else:
                st.error("❌ لا توجد بيانات في جدول أسماء الرسم")
                st.stop()
        
        # ============================================================
        # 2. اختيار الفترة
        # ============================================================
        
        periods_df = run_query('SELECT namee, no FROM "الفترة" ORDER BY no')
        period_names = periods_df['namee'].tolist() if periods_df is not None and not periods_df.empty else []
        
        if not period_names:
            st.error("❌ لا توجد فترات في جدول الفترة")
            st.stop()
        
        col_p1, col_p2, col_p3 = st.columns(3)
        with col_p1:
            start_p = st.selectbox("📅 من فترة:", period_names, key="catalog_start_period")
        with col_p2:
            end_p = st.selectbox("📅 إلى فترة:", period_names, index=len(period_names)-1, key="catalog_end_period")
        with col_p3:
            year = st.number_input("📅 العام:", min_value=2024, max_value=2030, value=2026)
        
        start_idx = period_names.index(start_p)
        end_idx = period_names.index(end_p)
        selected_periods = period_names[start_idx:end_idx+1]
        
        st.info(f"📅 الفترة المحددة: من {start_p} إلى {end_p}")
        
        # ============================================================
        # 3. جلب الأعمدة المتاحة
        # ============================================================
        
        all_columns = run_query('''
            SELECT "رقم اللوحة", "اسم العمود" as "الموقع", "العدد", "الشبكة", "الحجم", "المحافظة"
            FROM "اعمدة انارة" 
            WHERE "المحافظة" = %s AND "الحجم" = %s
        ''', (selected_city, selected_size))
        
        if all_columns is None or all_columns.empty:
            st.warning("⚠️ لا توجد أعمدة في هذه المحافظة والحجم")
            st.stop()
        
        # تحديد الأعمدة المحجوزة
        period_placeholders = ','.join([f"'{p}'" for p in selected_periods])
        booked_query = f'''
            SELECT DISTINCT "رقم اللوحة" FROM "حجوزات1" 
            WHERE "العام" = %s 
            AND "فترة الحجز" IN ({period_placeholders})
        '''
        booked_df = run_query(booked_query, (year,))
        booked_boards = booked_df['رقم اللوحة'].tolist() if booked_df is not None and not booked_df.empty else []
        
        # تصفية الأعمدة المتاحة
        available_columns = all_columns[~all_columns['رقم اللوحة'].isin(booked_boards)]
        
        if available_columns.empty:
            st.warning("⚠️ لا توجد لوحات متاحة في هذه المحافظة والحجم للفترة المحددة")
        else:
            # ============================================================
            # 4. عرض الإحصائيات
            # ============================================================
            
            st.subheader("📊 إحصائيات اللوحات المتاحة")
            
            col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
            with col_stat1:
                st.metric("إجمالي اللوحات", len(available_columns))
            with col_stat2:
                total_boards = available_columns['العدد'].sum()
                st.metric("إجمالي الوحدات", int(total_boards))
            with col_stat3:
                networks_count = available_columns['الشبكة'].nunique()
                st.metric("عدد الشبكات", networks_count)
            with col_stat4:
                st.metric("المحافظة", selected_city)
            
            # ============================================================
            # 5. عرض ملخص الشبكات
            # ============================================================
            
            st.subheader("📡 اللوحات المتاحة حسب الشبكة")
            
            network_summary = available_columns.groupby('الشبكة').agg({
                'رقم اللوحة': 'count',
                'العدد': 'sum'
            }).reset_index()
            network_summary.columns = ['الشبكة', 'عدد اللوحات', 'إجمالي الوحدات']
            
            st.dataframe(
                network_summary,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "الشبكة": "رقم الشبكة",
                    "عدد اللوحات": st.column_config.NumberColumn("عدد اللوحات"),
                    "إجمالي الوحدات": st.column_config.NumberColumn("إجمالي الوحدات")
                }
            )
            
            # ============================================================
            # 6. زر إضافة جميع الشبكات
            # ============================================================
            
            col_add_all1, col_add_all2, col_add_all3 = st.columns([1, 2, 1])
            with col_add_all2:
                if st.button("📦 إضافة جميع الشبكات للكتالوج", use_container_width=True, type="primary"):
                    all_board_numbers = available_columns['رقم اللوحة'].tolist()
                    
                    if selected_city not in st.session_state.catalog_selected_boards:
                        st.session_state.catalog_selected_boards[selected_city] = []
                    
                    for board in all_board_numbers:
                        if board not in st.session_state.catalog_selected_boards[selected_city]:
                            st.session_state.catalog_selected_boards[selected_city].append(board)
                    
                    if selected_city not in st.session_state.catalog_cities_added:
                        st.session_state.catalog_cities_added.append(selected_city)
                    
                    st.success(f"✅ تم إضافة جميع الشبكات ({len(all_board_numbers)} لوحة)")
            
            # ============================================================
            # 7. عرض اللوحات المختارة
            # ============================================================
            
            if st.session_state.catalog_selected_boards:
                st.divider()
                st.subheader("📋 اللوحات المختارة للكتالوج")
                
                total_selected = 0
                for city, boards in st.session_state.catalog_selected_boards.items():
                    if boards:
                        st.markdown(f"**📍 محافظة {city}:** {len(boards)} لوحة")
                        
                        board_placeholders = ','.join([f"'{b}'" for b in boards])
                        boards_details = run_query(f'''
                            SELECT "رقم اللوحة", "اسم العمود" as "الموقع", "العدد", "الشبكة"
                            FROM "اعمدة انارة" 
                            WHERE "رقم اللوحة" IN ({board_placeholders})
                            ORDER BY "الشبكة", "رقم اللوحة"
                        ''')
                        
                        if boards_details is not None and not boards_details.empty:
                            st.dataframe(
                                boards_details,
                                use_container_width=True,
                                hide_index=True,
                                column_config={
                                    "رقم اللوحة": "رقم اللوحة",
                                    "الموقع": "اسم الموقع",
                                    "العدد": st.column_config.NumberColumn("العدد"),
                                    "الشبكة": "رقم الشبكة"
                                }
                            )
                            
                            st.markdown("**🗑️ حذف من الكتالوج:**")
                            
                            col_remove1, col_remove2 = st.columns(2)
                            
                            with col_remove1:
                                networks_in_city = boards_details['الشبكة'].unique().tolist()
                                if networks_in_city:
                                    network_to_remove = st.selectbox(
                                        f"اختر شبكة لحذفها من {city}:",
                                        ["اختر الشبكة"] + networks_in_city,
                                        key=f"remove_network_{city}"
                                    )
                                    
                                    if network_to_remove != "اختر الشبكة" and st.button(f"🗑️ حذف شبكة {network_to_remove}", key=f"remove_network_btn_{city}"):
                                        network_boards = boards_details[boards_details['الشبكة'] == network_to_remove]['رقم اللوحة'].tolist()
                                        for board in network_boards:
                                            if board in st.session_state.catalog_selected_boards[city]:
                                                st.session_state.catalog_selected_boards[city].remove(board)
                                        st.success(f"✅ تم حذف شبكة {network_to_remove}")
                            
                            with col_remove2:
                                board_to_remove = st.selectbox(
                                    f"اختر لوحة لحذفها من {city}:",
                                    boards_details['رقم اللوحة'].tolist(),
                                    key=f"select_remove_{city}"
                                )
                                
                                if board_to_remove and st.button(f"🗑️ حذف اللوحة {board_to_remove}", key=f"remove_board_{city}"):
                                    if board_to_remove in st.session_state.catalog_selected_boards[city]:
                                        st.session_state.catalog_selected_boards[city].remove(board_to_remove)
                                    st.success(f"✅ تم حذف اللوحة {board_to_remove}")
                            
                            total_selected += len(boards)
                
                st.info(f"📊 إجمالي اللوحات المختارة: {total_selected}")
                
                col_controls1, col_controls2, col_controls3 = st.columns(3)
                with col_controls1:
                    if st.button("🔄 تحديث القائمة", use_container_width=True):
                        st.rerun()
                with col_controls2:
                    if st.button("🗑️ مسح جميع اللوحات", use_container_width=True):
                        st.session_state.catalog_selected_boards = {}
                        st.session_state.catalog_cities_added = []
                        st.success("✅ تم مسح جميع اللوحات")
                with col_controls3:
                    if total_selected > 0:
                        if st.button("📄 تصدير الكتالوج", use_container_width=True, type="primary"):
                            st.session_state.export_catalog = True
            else:
                st.info("📭 لم يتم اختيار أي لوحات بعد - استخدم زر 'إضافة جميع الشبكات'")
            
            # ============================================================
            # 8. تصدير الكتالوج (باستخدام دوال عرض السعر)
            # ============================================================
            
            if st.session_state.get('export_catalog', False) and st.session_state.catalog_selected_boards:
                st.divider()
                st.subheader("📤 تصدير الكتالوج")
                
                all_selected_boards = []
                for city, boards in st.session_state.catalog_selected_boards.items():
                    for board in boards:
                        all_selected_boards.append(board)
                
                if all_selected_boards:
                    board_placeholders = ','.join([f"'{b}'" for b in all_selected_boards])
                    all_boards_details = run_query(f'''
                        SELECT "رقم اللوحة", "اسم العمود" as "الموقع", "العدد", "الشبكة", "المحافظة"
                        FROM "اعمدة انارة" 
                        WHERE "رقم اللوحة" IN ({board_placeholders})
                        ORDER BY "المحافظة", "الشبكة", "رقم اللوحة"
                    ''')
                    
                    if all_boards_details is not None and not all_boards_details.empty:
                        col_exp1, col_exp2, col_exp3 = st.columns([1, 2, 1])
                        with col_exp2:
                            if st.button("📥 تحميل الكتالوج (Word)", use_container_width=True):
                                try:
                                    with st.spinner("جاري إنشاء الكتالوج..."):
                                        from docx import Document
                                        from docx.shared import Inches, Pt, RGBColor, Cm
                                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                                        from docx.oxml import OxmlElement
                                        from docx.oxml.ns import qn
                                        import io
                                        from datetime import datetime
                                        import os
                                        
                                        # ============================================================
                                        # دوال التنسيق المستخدمة في عرض السعر
                                        # ============================================================
                                        def _force_rtl_style(p):
                                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                            pPr = p._element.get_or_add_pPr()
                                            bidi = OxmlElement('w:bidi')
                                            bidi.set(qn('w:val'), '1')
                                            pPr.append(bidi)
                                            for run in p.runs:
                                                rPr = run._element.get_or_add_rPr()
                                                rtl = OxmlElement('w:rtl')
                                                rtl.set(qn('w:val'), '1')
                                                rPr.append(rtl)
                                        
                                        def set_table_rtl(table):
                                            tblPr = table._element.xpath('w:tblPr')[0]
                                            bidi = OxmlElement('w:bidiVisual')
                                            tblPr.append(bidi)
                                        
                                        # استخدام القالب
                                        template_path = 'template.docx'
                                        if os.path.exists(template_path):
                                            doc = Document(template_path)
                                        else:
                                            doc = Document()
                                        
                                        PURPLE_COLOR = "660099"
                                        
                                        # التاريخ
                                        today_date = datetime.now().strftime("%d / %m / %Y")
                                        p_date = doc.add_paragraph()
                                        p_date.add_run(f"التاريخ: {today_date}")
                                        _force_rtl_style(p_date)
                                        
                                        doc.add_paragraph()
                                        
                                        # العنوان
                                        p_title = doc.add_paragraph()
                                        p_title.add_run(f"كتالوج اللوحات الإعلانية المتاحة").bold = True
                                        _force_rtl_style(p_title)
                                        
                                        # مقدمة
                                        p_intro = doc.add_paragraph()
                                        p_intro.add_run(f"نقدم لكم اللوحات المتاحة من فترة ({start_p}) ولغاية ({end_p})")
                                        _force_rtl_style(p_intro)
                                        
                                        # عرض حسب المحافظة
                                        for city in all_boards_details['المحافظة'].unique():
                                            city_df = all_boards_details[all_boards_details['المحافظة'] == city]
                                            
                                            p_city = doc.add_paragraph()
                                            p_city.add_run(f"محافظة {city}").bold = True
                                            _force_rtl_style(p_city)
                                            
                                            p_size = doc.add_paragraph()
                                            p_size.add_run(f"لوحات قياس {selected_size}")
                                            _force_rtl_style(p_size)
                                            
                                            # عرض حسب الشبكة
                                            for network in city_df['الشبكة'].unique():
                                                network_df = city_df[city_df['الشبكة'] == network]
                                                
                                                # ============================================================
                                                # جدول 4 أعمدة بنفس تنسيق عرض السعر
                                                # ============================================================
                                                table = doc.add_table(rows=1, cols=4)
                                                table.style = 'Table Grid'
                                                set_table_rtl(table)
                                                
                                                for cell in table.columns:
                                                    cell.width = Cm(4.5)
                                                
                                                # رأس الجدول
                                                hdr = table.rows[0].cells
                                                hdr[0].text = "العدد"
                                                hdr[1].text = "رقم الشبكة"
                                                hdr[2].text = "العدد"
                                                hdr[3].text = "رقم الشبكة"
                                                
                                                for cell in hdr:
                                                    for p in cell.paragraphs:
                                                        if p.runs:
                                                            p.runs[0].bold = True
                                                            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                                                    tc_pr = cell._element.get_or_add_tcPr()
                                                    shd = OxmlElement('w:shd')
                                                    shd.set(qn('w:fill'), PURPLE_COLOR)
                                                    tc_pr.append(shd)
                                                    _force_rtl_style(p)
                                                
                                                # ملء الجدول بالبيانات
                                                rows_data = network_df.to_dict('records')
                                                total_units = int(network_df['العدد'].sum())
                                                
                                                # تجميع البيانات في أزواج
                                                paired_data = []
                                                for i in range(0, len(rows_data), 2):
                                                    if i + 1 < len(rows_data):
                                                        paired_data.append((rows_data[i], rows_data[i+1]))
                                                    else:
                                                        paired_data.append((rows_data[i], None))
                                                
                                                # الصف الأول: عنوان الشبكة والعدد الإجمالي
                                                first_row = table.add_row().cells
                                                first_row[0].text = str(total_units)
                                                first_row[1].text = f"الشبكة رقم {network}"
                                                first_row[2].text = ""
                                                first_row[3].text = ""
                                                
                                                for cell in first_row:
                                                    for p in cell.paragraphs:
                                                        _force_rtl_style(p)
                                                        if p.runs:
                                                            p.runs[0].bold = True
                                                
                                                # الصفوف المتبقية
                                                for pair in paired_data:
                                                    row_cells = table.add_row().cells
                                                    
                                                    if pair[0]:
                                                        row_cells[0].text = str(pair[0]['العدد'])
                                                        row_cells[1].text = str(pair[0]['الموقع'])
                                                    else:
                                                        row_cells[0].text = ""
                                                        row_cells[1].text = ""
                                                    
                                                    if pair[1]:
                                                        row_cells[2].text = str(pair[1]['العدد'])
                                                        row_cells[3].text = str(pair[1]['الموقع'])
                                                    else:
                                                        row_cells[2].text = ""
                                                        row_cells[3].text = ""
                                                    
                                                    for cell in row_cells:
                                                        for p in cell.paragraphs:
                                                            _force_rtl_style(p)
                                                
                                                doc.add_paragraph()
                                        
                                        # الإجمالي النهائي
                                        doc.add_paragraph()
                                        total_all = int(all_boards_details['العدد'].sum())
                                        p_grand = doc.add_paragraph()
                                        run_g = p_grand.add_run(f"العدد الإجمالي: {total_all}")
                                        run_g.bold = True
                                        run_g.font.size = Pt(14)
                                        run_g.font.color.rgb = RGBColor(102, 0, 153)
                                        _force_rtl_style(p_grand)
                                        
                                        doc.add_paragraph()
                                        
                                        # ملاحظة
                                        p_note = doc.add_paragraph()
                                        run_note = p_note.add_run("• ملاحظة: هذه المواقع متاحة للفترة المحددة.")
                                        run_note.bold = True
                                        _force_rtl_style(p_note)
                                        
                                        # حفظ وتحميل
                                        target = io.BytesIO()
                                        doc.save(target)
                                        target.seek(0)
                                        
                                        st.success("✅ تم إنشاء الكتالوج بنجاح!")
                                        st.download_button(
                                            label="📥 تحميل الكتالوج",
                                            data=target,
                                            file_name=f"كتالوج_{datetime.now().strftime('%Y%m%d')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            use_container_width=True
                                        )
                                        
                                        st.session_state.export_catalog = False
                                        
                                except Exception as e:
                                    st.error(f"❌ حدث خطأ أثناء إنشاء الكتالوج: {str(e)}")
                                    st.exception(e)
            
            # ============================================================
            # 9. خيارات إضافية
            # ============================================================
            
            with st.expander("⚙️ خيارات متقدمة", expanded=False):
                st.caption("خيارات إضافية لعرض الكتالوج")
                
                if is_admin():
                    show_prices = st.checkbox("💰 إظهار الأسعار (للمديرين فقط)")
                    if show_prices:
                        try:
                            fee_print, fee_ads = get_fees(draw_df, selected_size, "عادي", False)
                            st.info(f"""
                            💰 **تفاصيل الأسعار (للمديرين فقط):**
                            - سعر الطباعة الثابت: {fee_print}$
                            - سعر العرض الشهري: {fee_ads}$
                            """)
                        except:
                            st.warning("⚠️ لا يمكن جلب الأسعار حالياً")
    
    except Exception as e:
        st.error(f"❌ حدث خطأ: {str(e)}")
        st.exception(e)
# ============================================================
# إغلاق الاتصال
# ============================================================

conn.close()
